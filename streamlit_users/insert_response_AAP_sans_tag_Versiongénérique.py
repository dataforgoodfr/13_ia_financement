# Préambule 
"""
IMPORTANT :Pour que ce code fonctionne, il faut lui fournir une liste de dictionnaires nommée List_QuestionsResponses de la forme ci-dessous 
=> à remplir dans la partie "main" en fin de ce code :
List_QuestionsResponses =[
  {
    "question": "question 1",
    "response": "response to question 1" },
  {
    "question": "question 2",
    "response": "response to question 2" },
...etc
{
    "question": "question N",
    "response": "response to question N" } 
]
Les questions sont celles identifiées par l'IA dans le document AAP (Appel à Projets)
 et les réponses sont celles fournies par l'IA en réponse aux questions.
"""
# Import des bibliothèques
import docx
from docx import Document
from docx.document import Document as DocxDocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor
from docx.shared import Pt
import re
from datetime import datetime 
import logging
import uuid 
import glob 
import os
from pathlib import Path

# Get the directory of the current script (e.g., app.py)
SCRIPT_DIR = Path(__file__).parent.resolve()

def TestFullSimilarity(doc_1, doc_2):
    """
    Calculate the full similarity between two texts.
    """
    from fuzzywuzzy import fuzz
    Partial_Ratio = fuzz.partial_ratio(doc_1, doc_2)
    The_full_Ratio = fuzz.ratio(doc_1, doc_2)
    Partial_W_Ratio = fuzz.WRatio(doc_1, doc_2)
    EffectiveRatio = The_full_Ratio * 0.60 + Partial_W_Ratio * 0.20 + Partial_Ratio * 0.20
    print(f'EffectiveRatio = {EffectiveRatio} - FW Ratio = {The_full_Ratio} - FW Partial Ratio : {Partial_Ratio} - FW Partial W Ratio : {Partial_W_Ratio}')
    return EffectiveRatio

def TestSimilarity(doc_1, doc_2):
    """
    Calculate the similarity between two texts.
    """
    from fuzzywuzzy import fuzz
    Partial_Ratio = fuzz.partial_ratio(doc_1, doc_2)
    FuzzyWuzzy_Ratio = fuzz.ratio(doc_1, doc_2)
    The_global_Ratio = FuzzyWuzzy_Ratio * 0.80 + Partial_Ratio * 0.20
    print(f'Global Ratio = FuzzyWuzzy Ratio *0.8 + FuzzyWuzzy Partial Ratio *0.2  = {The_global_Ratio} - FuzzyWuzzy Ratio : {FuzzyWuzzy_Ratio} - FuzzyWuzzy Partial Ratio : {Partial_Ratio}')
    return The_global_Ratio

# Fonction pour scanner les paragraphes et les tables 
def iter_block_items(parent):
    if isinstance(parent, DocxDocumentType):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)
    

# Fonction pour insérer du texte dans un paragraphe
def Insert_Text_Paragraph(block_item, TextStart, TextEnd):
    """
    Args:
        block_item : the paragraph in which we insert the text
        TextStart : the text to be inserted at the beginning of the paragraph
        TextEnd : the text to be inserted at the end of the paragraph
    """
    if block_item.runs == []:
        block_item.text = TextStart + block_item.text + TextEnd 
    else: 
        block_item.runs[0].text = block_item.runs[0].text.replace("", TextStart,1) 
        NbRuns = block_item.runs.__len__()
        block_item.runs[NbRuns-1].text = block_item.runs[NbRuns-1].text.replace(block_item.runs[NbRuns-1].text, block_item.runs[NbRuns-1].text + TextEnd,1)
        block_item.runs[NbRuns-1].bold = False
        block_item.runs[NbRuns-1].italic = False
        block_item.runs[NbRuns-1].underline = False

    return


# Fonction pour insérer du texte dans un tableau 
def Insert_Text_Cell (tableCell, TextStart, TextEnd):
    """
    Args:
        tableCell : the cell of a table in which we insert the text
        TextStart : the tag to be inserted at the beginning of the paragraph
        TextEnd : the tag to be inserted at the end of the paragraph
    """
    ListOfRuns = []
    for paragCell in tableCell.paragraphs:
        ListOfRuns.extend(paragCell.runs)
    NbRuns = len(ListOfRuns)
    
    if NbRuns == 0:
        tableCell.text = tableCell.text.replace(tableCell.text, TextStart + tableCell.text + TextEnd)
    else: 
        ListOfRuns[0].text = ListOfRuns[0].text.replace("", TextStart,1) 
        ListOfRuns[NbRuns-1].text = ListOfRuns[NbRuns-1].text.replace(ListOfRuns[NbRuns-1].text, ListOfRuns[NbRuns-1].text + TextEnd,1)
        ListOfRuns[NbRuns-1].bold = False
        ListOfRuns[NbRuns-1].italic = False
        ListOfRuns[NbRuns-1].underline = False
    return


# Fonction pour voir si un paragraphe est le début d'une question
def Is_Para_Debut_Question (Text_Para, TheQuestion, SimilarityLimitStart):
    """
    Check if the paragraph is the beginning of a question.
    
    Args:
        Text_Para: The text of the paragraph.
        TheQuestion: The question text to compare against.
        SimilarityLimitStart: The similarity threshold to consider it as the start of a question.
    
    Returns:
        bool DebutQuestion: True if it is the start of a question, False otherwise.
        message: message explaining the result.
        bool FinQuestion: True if it is also the end of a question, False otherwise.
    """
    TheQuestion = str(TheQuestion.strip()) # on enlève les espaces au début et à la fin de la question
    Text_Para = str(Text_Para.strip()) # on enlève les espaces au début et à la fin de la question
    if TheQuestion == Text_Para and TheQuestion.strip() != '' and Text_Para.strip() != '':
        message = f"DebQuest.OK.1 : Début et fin de question trouvés en valeur exacte avec Text_Para_SSponct lu dans l'AAP : \n {Text_Para} \n pour la question :\n {TheQuestion}"
        DebutQuestion = True
        FinQuestion = True
        return DebutQuestion, message, FinQuestion 
    elif TheQuestion[:len(Text_Para)] == Text_Para and Text_Para.strip() != '' and TheQuestion.strip() != '':
        DebutQuestion = True
        FinQuestion = False
        message = f"DebQuest.OK.2 : Début de question trouvé en valeur exacte avec Text_Para_SSponct lu dans l'AAP : \n {Text_Para} \n pour la question :\n {TheQuestion}"
        return DebutQuestion, message, FinQuestion 
    elif TheQuestion == Text_Para[:len(TheQuestion)] and Text_Para.strip() != '' and TheQuestion.strip() != '':
        DebutQuestion = True
        FinQuestion = True
        message = f"DebQuest.OK.3 : Début de question trouvé en valeur exacte avec Text_Para_SSponct lu dans l'AAP : \n {Text_Para} \n pour la question :\n {TheQuestion}"
        return DebutQuestion, message, FinQuestion 
    elif TestSimilarity(TheQuestion[:len(Text_Para)], Text_Para) >= SimilarityLimitStart and Text_Para.strip() != '' and TheQuestion.strip() != '':
        DebutQuestion = True
        if (len(Text_Para) / len(TheQuestion) >0.9 and len(TheQuestion) / len(Text_Para) <1.1) or (len(TheQuestion) / len(Text_Para) >0.9 and len(Text_Para) / len(TheQuestion) <1.1)   :
            FinQuestion = True
            message = f"DebQuest.OK.4 : Début et fin de question trouvés par similarité avec Text_Para_SSponct lu dans l'AAP : \n {Text_Para} \n pour la question :\n {TheQuestion}"
        else:
            FinQuestion = False
            message = f"DebQuest.OK.5 : Début de question trouvé par similarité avec Text_Para_SSponct lu dans l'AAP : \n {Text_Para} \n pour la question :\n {TheQuestion}"
        return DebutQuestion, message, FinQuestion 
    else:
        DebutQuestion = False
        FinQuestion = False
        message = f"DebQuest.KO.6 : On n'a pas encore trouvé le début de la question, on est à Text_Para_SSponct: \n {Text_Para} \n pour la question :\n {TheQuestion}"
        return DebutQuestion, message, FinQuestion 

# Fonction pour voir si un texte correspond à une question
def Is_Text_Full_Question (Textlu, TheQuestion, SimilarityLimit, SimilarityMini, Last_Similarity=0):
    """
    Check if a text matches a question.
    
    Args:
        Textlu: The text read.
        TheQuestion: The question text to compare against.
        SimilarityLimit: The similarity threshold to consider it as the end of a question.
        SimilarityMini: The minimum similarity to consider it as the end of a question. Si on atteint un max de similarité (inflexion) sans atteindre ce mini, on annule toute la question.(remise de )
        Last_Similarity: The last similarity value found with the previous Textlu (default is 0).
    
    Returns:
        bool: True if it is the end of a question, False otherwise.
        message: message explaining the result.
        bool: True if the question is valid, False if we discover that it was finally not the right question (SimilarityMini not reached).
    """
    TheQuestion = str(TheQuestion.strip()) # on enlève les espaces au début et à la fin de la question
    Textlu = str(Textlu.strip()) # on enlève les espaces au début et à la fin de la question
    Similarite_Trouvee = TestFullSimilarity(Textlu, TheQuestion)
    if Textlu == TheQuestion:
        message = f"FullQuest.OK.1 : Fin de question trouvée en valeur exacte u dans l'AAP : \n {Textlu} \n pour la question :\n {TheQuestion}"
        return True, message, True # 1er True = fin de question - 2e True = question valide donc on OK pour insérer réponse
    elif TheQuestion[:len(Textlu)] == Textlu and Textlu.strip() != '':
        message = f"FullQuest.KO.2 : Correspondance partielle en valeur exacte - On n'est pas à la fin de la question, on est à Text_Para_SSponct: \n {Textlu} \n pour la question :\n {TheQuestion}"
        return False, message, True # 1er False = pas encore fin de quedtion - 2e True = question valide donc OK pour continuer à chercher la fin de question
    elif Similarite_Trouvee >= SimilarityLimit:
          message = f"FullQuest.OK.3 : Fin de question trouvée par similarité (dépassement de SimilarityLimit) avec texte lu dans l'AAP : \n {Textlu} \n pour la question :\n {TheQuestion}"
          return True, message, True # 1er True = fin de question - 2e True = question valide donc on OK pour insérer réponse
    elif Similarite_Trouvee < Last_Similarity and Last_Similarity >= SimilarityMini:
          message = f"FullQuest.OK.4 : Fin de question trouvée par similarité (inflexion du maximum) en dépassant le mini avec texte lu dans l'AAP : \n {Textlu} \n pour la question :\n {TheQuestion}"
          return True, message, True # 1er True = fin de question - 2e True = question valide donc on OK pour insérer réponse
    elif Similarite_Trouvee < Last_Similarity and Last_Similarity - Similarite_Trouvee > 7  and Last_Similarity < SimilarityMini: 
          message = f"FullQuest.K0.5 : Inflexion conséquente (+ de 7) de maximum de similarité atteinte sans dépasser la similarité mini => le début de question qui avait été trouvé était erroné et on annule tout avec texte lu dans l'AAP : \n {Textlu} \n pour la question :\n {TheQuestion}"
          return False, message, False # 1er False = pas fin de question - 2e False = question non valide (car inflexion maxi sans atteindre le mini donc il faut tout annuler        
    else:
        message = f"FullQuest.KO.6 : On n'est pas à la fin de la question, on est à Text_Para_SSponct: \n {Textlu} \n pour la question :\n {TheQuestion}"
        return False, message, True # 1er False = pas fin de question - 2e True = question valide donc OK pour continuer à chercher la fin de question


# Fonction pour lire les questions et la taille des réponses souhaitée
def Write_Answers_in_docx (List_QuestionsResponses, PathFolderSource, PathForOutputsAndLogs):
    """
    Insère les réponses dans les fichiers Word.
    Génère un fichier Word annoté + un fichier récapitulatif Q/A.

    Args:
        List_QuestionsResponses : list of questions and answers
        PathFolderSource: Path to the folder containing the files to be read
        PathForOutputsAndLogs: Path to the folder where we put the documents with answers and containing the log file
    """
    import re
    SimilarityLimit = 95 # Seuil de similarité qui déclenche l'insertion d'une réponse
    SimilarityLimitStart = 95
    SimilarityMini =50 # Minimum de similarité pour considérer qu'on a trouvé la fin de la question (pour les cas où la similarité atteint un max mais n'est pas assez élevée au dessus de ce mini)
    EverythingOK = True 
    FilesWithPath = []
    Go_InsertResponse = False 
    Question_LuDoc = '' # texte de la question lue dans l'AAP reconstitué si nécessaire par concaténation de plusieurs paragraphes
    Question_LuDoc_SSponct = '' # texte de la question lue dans l'AAP reconstitué ci-avant dont on a retiré la ponctuation (espaces, points, virgules,..)
    Question_LuIA = '' # texte de la question lue par l'IA et fournie dans un docitionnaire
    Question_LuIA_SSponct = '' # texte de la question lue par l'IA ci-avant dont on a retiré la ponctuation (espaces, points, virgules,..)
    Text_Para = '' # texte du paragraphe courant lu dans l'AAP 
    Text_Para_SSponct = '' # texte du paragraphe lu dans l'AAP dont on a retiré la ponctuation (espaces, points, virgules,..)
    FlagDébutQuestion = False # Flag pour savoir si on a trouvé le début de la question
    FlagFinQuestion = False # Flag pour savoir si on a trouvé la fin de la question
    NumBoucle = 0
    NB_Rows = 1 # Nombre de lignes dans le tableau par défaut
    NB_Cols = 1 # Nombre de colonnes dans le tableau par défaut
    # Récupération des fichiers avec chemins absolus
    FilesWithPath = [file.resolve() for file in PathFolderSource.glob('*.*')]

    # Affichage des résultats (pour vérification)
    print(FilesWithPath)

    #===============================================

    for file in FilesWithPath:
        TheExtension = file.suffix
        if TheExtension != ".docx":
            EverythingOK = False
            MessageError = str(datetime.now()) + ' Error encountered when reading files : There should only be docx files in the folder'
            logging.error(MessageError)
            print (MessageError)
        else: # Docx OK 
            EverythingOK = True
            NameOfDocument = file.name.split('.')[0]
            if EverythingOK : 
                try:
                    docWork = docx.Document(file)
                except IOError:
                    MessageError = str(datetime.now()) + ' Error encountered when reading Word docx file ' + file
                    logging.error(MessageError)
                    print(MessageError)

                # Traitement de chaque question fournie par l'IA
                for value in List_QuestionsResponses:
                    # Nettoyage du texte de la question fournie par l'IA pour enlever la ponctuation, les tabulations et les retours à la ligne
                    Question_LuIA = str(value ["question"])
                    Question_LuIA_SSponct = str(re.sub(r'[ ,.\'\n’‘□☐:();_–  °]','', Question_LuIA)) # cas idéal : question et paragraphe identiques en enlevant la ponctuation
                    TheResponse = value ["response"]

                    # Initialisation des variables
                    Question_LuDoc = '' # texte de la question lue dans l'AAP reconstitué si nécessaire par concaténation de plusieurs paragraphes
                    Question_LuDoc_SSponct = '' # texte de la question lue dans l'AAP reconstitué ci-avant dont on a retiré la ponctuation (espaces, points, virgules,..)
                    Text_Para = '' # texte du paragraphe courant lu dans l'AAP 
                    Text_Para_SSponct = '' # texte du paragraphe lu dans l'AAP dont on a retiré la ponctuation (espaces, points, virgules,..)
                    FlagDébutQuestion = False # Flag pour savoir si on a trouvé le début de la question
                    FlagFinQuestion = False
                    Go_InsertResponse = False # Flag pour savoir si on a trouvé la question complète et qu'on peut insérer la réponse
                    Precedente_Similarite_Trouvee = 0 # Similarité précédente trouvée par la fonction Is_Text_Full_Question  

                    # Boucle sur les paragraphes et les tableaux du document
                    for block_item in iter_block_items(docWork): 
                    # ============================================
                    # 1 - Cas paragraphe plein text (pas tableau)
                    # ============================================
                        if isinstance(block_item, Paragraph): 
                            if block_item.text.strip() != '':
                    # envoi au terminal d'un message d'avancement
                                NumBoucle += 1
                                print()
                                print(f'Boucle n° {NumBoucle} sur le paragraphe : {block_item.text}')

                    # Nettoyage du texte du paragraphe pour enlever la ponctuation, les tabulations et les retours à la ligne
                                Text_Para = str(block_item.text)
                                Text_Para = Text_Para.replace("\t", "") # suppression des tabulations
                                Text_Para = Text_Para.replace("\n", "") # suppression des retours à la ligne
                                Text_Para_SSponct = str(re.sub(r'[ ,.\'\n’‘□☐:();_–  °]','', Text_Para))
                                Text_Para_SSponct = Text_Para_SSponct.replace(" ", "") # suppression des espaces restants ayant malheureusement échappé aux effacements précédents

                    # Si on est déjà dans une question, on concatène le texte du paragraphe au texte de la question et on évalue si on atteint la fin de la question
                                if FlagDébutQuestion == True:
                                    Precedente_Similarite_Trouvee = TestFullSimilarity(Question_LuDoc_SSponct, Question_LuIA_SSponct)
                                    Question_LuDoc = str(Question_LuDoc + Text_Para )
                                    Question_LuDoc_SSponct = str(Question_LuDoc_SSponct + Text_Para_SSponct)                         
                                    Go_InsertResponse, message, FlagDébutQuestion = Is_Text_Full_Question (Question_LuDoc_SSponct, Question_LuIA_SSponct, SimilarityLimit, SimilarityMini, Precedente_Similarite_Trouvee)
                                    # On reévalue ici à chaque fois FlagDébutQuestion : Si la question est valide il reste à True, sinon il repasse à False (question finalement invalide car ayant atteint le maxi de similarité sans atteinte le seuil mini de similarité)
                                    print(message)
                                    if Go_InsertResponse and FlagDébutQuestion:
                                        # on a trouvé la question complète dans le document et on insère la réponse
                                        Insert_Text_Paragraph(block_item, "" , "\n" + TheResponse)
                                        # Réinitialisation des variables puisqu'on a inséré la réponse
                                        Question_LuDoc = '' 
                                        Question_LuDoc_SSponct = ''
                                        Question_LuIA = '' 
                                        Question_LuIA_SSponct = '' 
                                        Text_Para = '' 
                                        Text_Para_SSponct = '' 
                                        FlagDébutQuestion = False
                                        FlagFinQuestion = False
                                        Go_InsertResponse = False
                                        Precedente_Similarite_Trouvee = 0 # Similarité précédente trouvée par la fonction Is_Text_Full_Question  
                                        break # on sort de la boucle car on a inséré la réponse et on passe à la question IA suivante
                                    else:
                                        pass
                                        # On est au milieu de la question côté document mais on n'a pas le go pour insérer la réponse

                    # Si on n'est pas encore dans une question, on compare le texte avec la question pour détecter un début de question
                                else: # FlagDébutQuestion == False  => on n'est pas encore dans la question dans le document
                                    # comparer le texte du paragraphe sans ponctuation avec le texte de la question sans ponctuation
                                    # et si OK, le retour de la fonction fait passer FlagDébutQuestion = True
                                    FlagDébutQuestion, message, FlagFinQuestion = Is_Para_Debut_Question(Text_Para_SSponct, Question_LuIA_SSponct, SimilarityLimitStart)
                                    print(message)

                        # Si on a trouvé du même coup à la fois le début et la fin de question, on insère la réponse
                                    if FlagDébutQuestion and FlagFinQuestion: #on a trouvé dans le doc le début et la fin de la question
                                        # insérer la réponse dans le paragraphe
                                        Insert_Text_Paragraph(block_item, "" , "\n" + TheResponse)
                                        # Réinitialisation des variables puisqu'on a inséré la réponse
                                        Question_LuDoc = '' 
                                        Question_LuDoc_SSponct = ''
                                        Question_LuIA = '' 
                                        Question_LuIA_SSponct = '' 
                                        Text_Para = '' 
                                        Text_Para_SSponct = '' 
                                        FlagDébutQuestion = False
                                        FlagFinQuestion = False
                                        Go_InsertResponse = False
                                        Precedente_Similarite_Trouvee = 0 # Similarité précédente trouvée par la fonction Is_Text_Full_Question  
                                        break # on sort de la boucle car on a inséré la réponse et on passe à la question IA suivante

                        # Si on a trouvé le début mais pas la fin de question, on initialise le texte du début de question
                                    elif FlagDébutQuestion and not(FlagFinQuestion): #on a trouvé dans le doc le début de la question mais pas la fin de la question
                                        Question_LuDoc = str(Text_Para)
                                        Question_LuDoc_SSponct = str(Text_Para_SSponct)

                            else:
                                pass # c'est un paragraphe vide donc on l'ignore

                    # ============================================
                    # 2 - Cas paragraphe de tableau (cellule)
                    # ============================================

                        elif isinstance(block_item, Table): 
                            # creation d'une liste des textes des cellules du tableau avec Num Ligne, Num colonne et texte de la cellule et supprimant les doublons des lignes et colonnes fusionnées
                            NB_Rows = len(block_item.rows)
                            NB_Cols = len(block_item.columns)
                            List_Text_Cell = []
                            DictCell = {}
                            for row in range(NB_Rows): 
                                for col in range(NB_Cols): 
                                    if col!=0 and str(block_item.cell(row, col).text) == str(block_item.cell(row, col-1).text):# élimination des doublons des lignes fusionnées = cas où les 2 colonnes de la ligne ont le même contenu
                                        pass
                                    elif row!=0 and str(block_item.cell(row, col).text) == str(block_item.cell(row-1, col).text):# élimination des doublons des colonnes fusionnées = cas où les 2 lignes de la colonne ont le même contenu
                                        pass
                                    else: # on ajoute le texte de la cellule dans la liste des textes du tableau
                                        DictCell ["col"] = col
                                        DictCell ["row"] = row
                                        DictCell ["textcell"] = str(block_item.cell(row, col).text)
                                        new_dict = DictCell.copy() 
                                        List_Text_Cell.append ( new_dict ) 
                                        DictCell.clear() 
                            # à ce stade, le dictionnaires des textes du tableau a été purgé des doublons liés aux cellules fusionnées et on peut comparer avec la question
                            # boucle sur les cellules du tableau
                            for dict_item in List_Text_Cell:
                                row = dict_item["row"]
                                col = dict_item["col"]
                                Text_Para = dict_item ["textcell"]
                                if Text_Para.strip() != '':
                                    

                                # envoi au terminal d'un message d'avancement
                                    NumBoucle += 1
                                    print()
                                    print(f'Boucle n° {NumBoucle} sur la cellule : {Text_Para}')

                                # Nettoyage du texte du paragraphe pour enlever la ponctuation, les tabulations et les retours à la ligne
                                    Text_Para = str(block_item.cell(row, col).text)
                                    Text_Para = Text_Para.replace("\t", "") # suppression des tabulations
                                    Text_Para = Text_Para.replace("\n", "") # suppression des retours à la ligne
                                    Text_Para_SSponct = str(re.sub(r'[ ,.\'\n’‘□☐:();_–  °]','', Text_Para))
                                    Text_Para_SSponct = Text_Para_SSponct.replace(" ", "") # suppression des espaces restants ayant malheureusement échappé aux effacements précédents

                                  # Si on est déjà dans une question, on concatène le texte du paragraphe au texte de la question et on évalue si on atteint la fin de la question
                                    if FlagDébutQuestion == True:
                                        Precedente_Similarite_Trouvee = TestFullSimilarity(Question_LuDoc_SSponct, Question_LuIA_SSponct)
                                        Question_LuDoc = str(Question_LuDoc + Text_Para)
                                        Question_LuDoc_SSponct = str(Question_LuDoc_SSponct + Text_Para_SSponct)                        
                                        Go_InsertResponse, message, FlagDébutQuestion = Is_Text_Full_Question (Question_LuDoc_SSponct, Question_LuIA_SSponct, SimilarityLimit, SimilarityMini, Precedente_Similarite_Trouvee)
                                        # On reévalue ici à chaque fois FlagDébutQuestion : Si la question est valide il reste à True, sinon il repasse à False (question finalement invalide car ayant atteint le maxi de similarité sans atteinte le seuil mini de similarité)
                                        print(message)
                                        if Go_InsertResponse and FlagDébutQuestion:
                                            # on a trouvé la question complète dans le document et on insère la réponse
                                            Insert_Text_Cell (block_item.cell(row, col), '' , '\n' + TheResponse)
                                            # Réinitialisation des variables puisqu'on a inséré la réponse
                                            Question_LuDoc = '' 
                                            Question_LuDoc_SSponct = ''
                                            Question_LuIA = '' 
                                            Question_LuIA_SSponct = '' 
                                            Text_Para = '' 
                                            Text_Para_SSponct = '' 
                                            FlagDébutQuestion = False
                                            FlagFinQuestion = False
                                            Go_InsertResponse = False
                                            Precedente_Similarite_Trouvee = 0 # Similarité précédente trouvée par la fonction Is_Text_Full_Question  
                                            break # on sort de la boucle car on a inséré la réponse et on passe à la question IA suivante
                                        else:
                                            pass
                                            # On est au milieu de la question côté document mais on n'a pas le go pour insérer la réponse


                                    # Si on n'est pas encore dans une question, on compare le texte avec la question pour détecter un début de question
                                    else: # FlagDébutQuestion == False  => on n'est pas encore dans la question dans le document
                                        # comparer le texte du paragraphe sans ponctuation avec le texte de la question sans ponctuation
                                        # et si OK, le retour de la fonction fait passer FlagDébutQuestion = True
                                        FlagDébutQuestion, message, FlagFinQuestion = Is_Para_Debut_Question(Text_Para_SSponct, Question_LuIA_SSponct, SimilarityLimitStart)
                                        print(message)

                                    # Si on a trouvé du même coup à la fois le début et la fin de question, on insère la réponse
                                        if FlagDébutQuestion and FlagFinQuestion: #on a trouvé dans le doc le début et la fin de la question
                                            # insérer la réponse dans le paragraphe
                                            Insert_Text_Cell (block_item.cell(row, col), '' , '\n' + TheResponse)
                                            # Réinitialisation des variables puisqu'on a inséré la réponse
                                            Question_LuDoc = '' 
                                            Question_LuDoc_SSponct = ''
                                            Question_LuIA = '' 
                                            Question_LuIA_SSponct = '' 
                                            Text_Para = '' 
                                            Text_Para_SSponct = '' 
                                            FlagDébutQuestion = False
                                            FlagFinQuestion = False
                                            Go_InsertResponse = False
                                            Precedente_Similarite_Trouvee = 0 # Similarité précédente trouvée par la fonction Is_Text_Full_Question  
                                            break # on sort de la boucle car on a inséré la réponse et on passe à la question IA suivante

                                    # Si on a trouvé le début mais pas la fin de question, on initialise le texte du début de question
                                        elif FlagDébutQuestion and not(FlagFinQuestion): #on a trouvé dans le doc le début de la question mais pas la fin de la question
                                            Question_LuDoc = str(Text_Para)
                                            Question_LuDoc_SSponct = str(Text_Para_SSponct)
                                else:
                                    pass # c'est un paragraphe vide donc on l'ignore
                                #List_Text_Cell.clear() # on vide la liste des cellules pour le prochain tableau

                        else:#cas ni plein texte ni tableau => inconnu
                            MessageError = str(datetime.now()) + ' erreur00 en lecture du docx : ni plein texte ni tableau => cas inconnu'
                            logging.error(MessageError)
                            print(MessageError)

                
                # === Sauvegarde du fichier final avec réponses
                timestamp = datetime.now().strftime("%Y-%m-%d_%Hh%Mmn%Ss")
                output_filename = f'{NameOfDocument}_with_answers_{timestamp}.docx'
                path_output_doc = PathForOutputsAndLogs / f'{output_filename}'
                docWork.save(path_output_doc)

                # === Génération du document Q&A séparé
                documentQA = Document()
                title = documentQA.add_heading(
                    f'Liste des questions/réponses pour :\n{NameOfDocument}\nHeure : {timestamp}\n', level=1
                )
                title.style.font.size = Pt(14)
                # Ligne de séparation
                documentQA.add_paragraph("             _________________________________________________")
                documentQA.add_paragraph().add_run().add_break()

                for value in List_QuestionsResponses:
                    p = documentQA.add_paragraph()
                    question_clean = value["question"]# .replace("??", "").replace(TagQStart, "").replace(TagQEnd, "")
                    run = p.add_run(question_clean)
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    documentQA.add_paragraph('\n' + value["response"] + '\n')

                qa_filename = f'{NameOfDocument}_Q-A_{timestamp}.docx'
                path_qa_doc = PathForOutputsAndLogs / f'{qa_filename}'

                # Sauvegarde du document
                documentQA.save(path_qa_doc)
                documentQA = Document()

                print('✅ Fin du programme d’écriture des réponses dans les fichiers.')
            else :#if EverythingOK pas OK
                print("FilesWithPath is not OK, please check the files in the folder")
                MessageError = str(datetime.now()) + ' Error encountered when reading Word docx file , please check type .docx and name of the file with no UID)'
                logging.error(MessageError)
                print(MessageError)

    return path_output_doc, path_qa_doc, output_filename, qa_filename



#@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@ MAIN PROGRAM @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@

# imports
from docx import Document # import de python-docx
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor

import re
from datetime import datetime # for log file
import logging # for log file
import uuid # for unique ID creation
import glob # for file opening & reading


# Below is the liste of Q/A given by the llm that must be inserted into the AAP. 
# It is empty here for the program to fill it with the output of the llm (integration phase to be done)
# The format is a list of dictionnaries with "question" and "response" keys
List_QuestionsResponses =[
  {
    "question": "",
    "response": "" },
  {
    "question": "",
    "response": "" }
]

# Get the directory of the current script (e.g., app.py)
SCRIPT_DIR = Path(__file__).parent.resolve()
# Settings for the path files
#Path_where_we_put_Outputs = r'/Users/jfm/Library/CloudStorage/OneDrive-Personnel/Python yc Dev D4G/3 - Dev IA Asso/Pour les logs/' 
#Folder_where_the_files_are = r'/Users/jfm/Library/CloudStorage/OneDrive-Personnel/Python yc Dev D4G/3 - Dev IA Asso/LesFilesA Lire/'
Path_where_we_put_Outputs = SCRIPT_DIR / "Pour les logs/" # répertoire pour l'AAP source (non rempli)
Folder_where_the_files_are = SCRIPT_DIR / "LesFilesA Lire/" # répertoire pour l'AAP source (non rempli)

# Création du dossier parent si nécessaire
Path(Path_where_we_put_Outputs).mkdir(parents=True, exist_ok=True)
Path(Folder_where_the_files_are).mkdir(parents=True, exist_ok=True)
#activate logging of errors in a txt file
logging.basicConfig(filename=str(Path_where_we_put_Outputs) + r'/logs-IA_for_Asso.txt')


# Write the answers into the docx files just below the questions
path_output_doc, path_qa_doc, output_filename, qa_filename = Write_Answers_in_docx(List_QuestionsResponses, Folder_where_the_files_are, Path_where_we_put_Outputs)
print(f"path vers Fichier avec réponses : {path_output_doc}")
print(f"path vers Fichier Q&A : {path_qa_doc}")
print(f"Nom du fichier avec réponses : {output_filename}")
print(f"Nom du fichier Q&A : {qa_filename}")
#@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@ END OF MAIN PROGRAM @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@