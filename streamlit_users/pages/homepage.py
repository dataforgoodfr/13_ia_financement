# Import personal functions
from rag_pipelines import process_new_doc, process_existing_doc, QA_pipeline, update_hybrid_rag_wrapper
from graphrag_retriever import load_knowledgeGraph_vis
from read_answer_aap import Read_Questions_in_docx, Write_Answers_in_docx

# import langchain libraries
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter

# import other libraries
from io import StringIO
import asyncio
import nest_asyncio
import os
import pandas as pd
import json
from pathlib import Path
import streamlit as st

# to download APP in docx format
from docx import Document
from docx.shared import Pt
from io import BytesIO

# Patch torch.classes pour éviter l'erreur de Streamlit
import torch

#=========corrections incompatibilités streamlit / event loop
#=============== nécessaire pour streaming des réponses graphRAG
# Patch the event loop to allow nested async calls
nest_asyncio.apply()
#==========fin


#=================== Correction conflit event loops torch/streamlit
# Save the original __getattr__ method
original_getattr = torch._classes._Classes.__getattr__

# Define a patched version to handle __path__
def patched_getattr(self, attr):
    if attr == "__path__":
        # Explicitly block access to __path__
        raise AttributeError(f"'{self.__class__.__name__}' has no attribute '__path__'")
    return original_getattr(self, attr)

# Apply the patch
torch._classes._Classes.__getattr__ = patched_getattr
#===============fin

# Get the directory of the current script (e.g., app.py)
# SL: add 1 more level parent because homepage is in pages folder
SCRIPT_DIR = Path(__file__).parent.parent.resolve()

#=================== Functions definition
#doc_category: the type of document, must be 'pp' or 'asso'

def stream_pathRAG_response(stream_resp, response_container):
    async def stream_response():
        response_buffer = StringIO()        
        
        # Get the existing event loop
        loop = asyncio.get_event_loop()
        
        # Process the async generator
        async for chunk in stream_resp["response_stream"]:
            response_buffer.write(chunk)
            response_container.markdown(response_buffer.getvalue())

        st.session_state["full_response"] = response_buffer.getvalue()
        response_buffer.close()


    # Run in Streamlit's existing event loop
    loop = asyncio.get_event_loop()
    loop.run_until_complete(stream_response())

def stream_hybridRAG_response(stream_resp, response_container):
    # 1. Créer un buffer pour accumuler la réponse
    response_buffer = StringIO()
    response_container = st.empty()  # Conteneur vide pour mise à jour dynamique

    for chunk in stream_resp["response_stream"]:
    # Ajouter le token au buffer
        response_buffer.write(chunk)
        # Mettre à jour le conteneur avec le contenu accumulé
        response_container.markdown(response_buffer.getvalue())

    # 3. Récupérer la réponse complète
    st.session_state["full_response"] = response_buffer.getvalue()
    response_buffer.close()

def generate_docx(data):
    doc = Document()
    
    # Style du titre
    title = doc.add_heading('Questions/Réponses', 0)
    title.style.font.size = Pt(14)
    
    for item in data:
        # Ajout de la question
        doc.add_heading(item['question'], level=1)
        
        # Ajout de la réponse
        doc.add_paragraph(item['response'])
        
        # Ligne de séparation
        doc.add_paragraph().add_run().add_break()
    
    # Sauvegarde en mémoire
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

# log loading of available documents in database
def load_doc_traces(doc_category: str):
    # load hashes and save them in a dataframe
    files_paths=[SCRIPT_DIR/ "hybridrag_hashes.json", SCRIPT_DIR/"graphrag_hashes.json"]
    table=pd.DataFrame([])

    for file_path in files_paths:
        exising_hashes={}
        if os.path.exists(file_path):                
            t=pd.read_json(file_path,).T
            table=pd.concat([table, t], axis=0)

    # préparer les hashes
    if len(table)>0:
        # enlever les lignes en double hybrid rag/ graph rag
        table=table.drop_duplicates(subset=["Nom du doc","Titre auto","Taille du texte (en car)"])
        table=table.sort_values("Date de création", ascending=False)
        table=table[table["doc_category"]==doc_category].drop(columns=["rag_type"])
        table['Date de création'] = pd.to_datetime(table['Date de création']).dt.floor('s')
        return table     
    else:
        return []

# load traces of pp and asso docs in the databases
df_existing_pp= load_doc_traces("pp")
df_existing_asso= load_doc_traces("asso")

#=================== UI v2 Streamlit app
# Font definition and global config off the streamlit app is in the config.toml file

# --- CSS styling ---

def load_css(file_path):
    with open(file_path) as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

css_path =SCRIPT_DIR/ "styles.css"
#css_path = pathlib.Path("styles.css")
load_css(css_path)


# --- Homepage config ---
st.set_page_config(
    page_title="DossierIA+",
    page_icon=":clipboard:",
    layout="wide",
    initial_sidebar_state="auto",
)

# --- Sidebar ---
st.sidebar.title("Mon Compte")
st.sidebar.image(SCRIPT_DIR/ "img/logo_D4G_no_text_back.png", width=100)
st.sidebar.markdown("**Bienvenue dans DossierIA+ !**")

# --- Main header ---
st.title("DossierIA+")
st.header("Facilitez vos réponses aux demandes de subventions")
st.write("""DossierIA+ pré-remplit automatiquement vos formulaires à partir de vos documents sources. **En quelques clics, votre dossier est prêt !**  
         Vous gardez le contrôle pour les ajuster facilement, tout en gagnant un temps précieux. Conçue par des bénévoles de Data For Good avec les associations du groupe SOS pour les associations, cette solution vous aide à maximiser l’impact de vos actions en vous permettant de vous concentrer sur l’essentiel.""")
st.divider() 

# --- 1st step Project uploading documents pp ---
st.header("Etape 1 - Chargez vos documents")
st.subheader("Votre projet")

# Notification zone if a document is uploaded in session state or not
zone_notif_doc_charge_pp=st.empty()
if "current_pp_in_use" in st.session_state:
    zone_notif_doc_charge_pp.success(f"""Document chargé: {st.session_state["current_pp_in_use"]}""")
else:
    zone_notif_doc_charge_pp.warning(f"Veuillez charger un document projet")

# Drag and drop multiple files UI
uploaded_project_files = st.file_uploader("**Chargez les documents présentant votre projet** (note de cadrage, dossier de présentation, budget, annexes, etc.) ", 
    accept_multiple_files=True, type=['pdf', 'docx', 'csv'], key='uploader_proj')

if uploaded_project_files is not None:
    st.session_state["uploaded_pp"]=uploaded_project_files
    for uploaded_file in uploaded_project_files:
        st.write(f"Fichier {uploaded_file.name} importé avec succès.")
                
# Naming, upload and process of pp file
col1_load_pp, col2_load_pp = st.columns([4, 1], vertical_alignment="center")
with col1_load_pp:
    proj_name = st.text_input("**Titre de votre projet**", placeholder="Ex. : Prévention des zoonoses au Cambodge", key='proj_name')

with col2_load_pp:
    upload_button = st.button("Importer", key='btn_proj')
    st.session_state["upload_button"]=upload_button
    if upload_button:
        if proj_name and uploaded_project_files:
            st.session_state["uploaded_pp"]=uploaded_project_files
            with st.spinner("Wait for it...", show_time=True):
                for message in process_new_doc(uploaded_project_files, proj_name, "pp"):
                    st.markdown(message, unsafe_allow_html=True)

            st.session_state["current_pp_in_use"]=proj_name
            st.success(f"✅ Projet '{proj_name}' importé avec succès.")
            with zone_notif_doc_charge_pp:
                    st.success(f"""Document chargé: {st.session_state[f"current_pp_in_use"]}""")
        else:
            st.warning("Veuillez saisir un titre de projet avant d'importer.")


# loading pp documents from the database
col3_existing_pp, col4_existing_pp = st.columns([4, 1], vertical_alignment="center")
with col3_existing_pp:
    #proj_lib = st.selectbox("Ou utilisez des documents de la bibliothèque", ["Présentation Projet Mahakam"], key='proj_lib')
    list_doc_names=[]
    if len(df_existing_pp)>0:
        list_doc_names=df_existing_pp[f"Nom du doc"].unique() 
        selected_doc_name= st.selectbox("Ou utilisez des documents de la bibliothèque", options=list_doc_names, key='selected_pp_lib')
        
with col4_existing_pp:
    upload_lib_button = st.button("Charger", key='btn_proj_lib')
    if upload_lib_button and selected_doc_name:
        hash=df_existing_pp[df_existing_pp[f"Nom du doc"]==selected_doc_name].index.values[0]
        st.session_state["selected_pp_name"]=selected_doc_name
        
        with st.spinner("Wait for it...", show_time=True):
                for message in process_existing_doc(hash, selected_doc_name, "pp"):
                    st.markdown(message, unsafe_allow_html=True)

        st.session_state["current_pp_in_use"]=selected_doc_name
        st.success(f"✅ Projet de la bibliothèque : {selected_doc_name} importé avec succès.")
        with zone_notif_doc_charge_pp:
                st.success(f"""Document chargé: {st.session_state["current_pp_in_use"]}""")
st.markdown("")   


# --- Association uploading documents ---
st.subheader("Votre association")

# Notification zone if a document is uploaded in session state or not
zone_notif_doc_charge_asso=st.empty()
if "current_asso_in_use" in st.session_state:
    zone_notif_doc_charge_asso.success(f"""Document chargé: {st.session_state["current_asso_in_use"]}""")
else:
    zone_notif_doc_charge_asso.warning(f"Veuillez charger un document association")


# Drag and drop multiple files
uploaded_asso_files = st.file_uploader("**Chargez les documents présentant votre association** (présentation institutionnelle, statuts, rapports d’activité, etc.) ", 
    accept_multiple_files=True, type=['pdf', 'docx', 'csv'], key='assoc')
if uploaded_asso_files is not None:
    st.session_state["uploaded_asso"]=uploaded_asso_files
    for uploaded_file in uploaded_asso_files:
        st.write(f"Fichier {uploaded_file.name} importé avec succès.")
                
# Naming, upload and process of asso file
col1_load_asso, col2_load_asso = st.columns([4, 1], vertical_alignment="center")
with col1_load_asso:
    assoc_name = st.text_input("**Nom associé aux fichiers**", placeholder="Ex : Présentation Forêts en Danger", key='assoc_name')
with col2_load_asso:
    upload_button_asso = st.button("Importer", key='btn_asso')
    st.session_state["upload_button_asso"]=upload_button_asso
    if upload_button_asso:
        if assoc_name and uploaded_asso_files:
            st.session_state["uploaded_asso"]=uploaded_asso_files
            with st.spinner("Wait for it...", show_time=True):
                    for message in process_new_doc(uploaded_asso_files, assoc_name, "asso"):
                        st.markdown(message, unsafe_allow_html=True)
            st.session_state["current_asso_in_use"]=assoc_name
            st.success(f"✅ Document '{assoc_name}' importé avec succès.")
            with zone_notif_doc_charge_asso:
                st.success(f"""Document chargé: {st.session_state[f"current_asso_in_use"]}""")
        else:
            st.warning("Veuillez saisir un titre de document avant d'importer.")

# Uploading asso documents from the library 
col3_existing_asso, col4_existing_asso = st.columns([4, 1], vertical_alignment="center")
with col3_existing_asso:
    #assoc_lib = st.selectbox("Ou utilisez des documents de la bibliothèque", ["Présentation Projet Mahakam"], key='assoc_lib')
    list_doc_asso_names=[]
    if len(df_existing_asso)>0:
        list_doc_asso_names=df_existing_asso[f"Nom du doc"].unique() 
        selected_doc_asso_name= st.selectbox("Ou utilisez des documents de la bibliothèque", options=list_doc_asso_names, key='selected_asso_lib')
        
with col4_existing_asso:
    upload_lib_asso_button = st.button("Charger", key='btn_assoc_lib')
    if upload_lib_asso_button and selected_doc_asso_name:
        hash=df_existing_asso[df_existing_asso[f"Nom du doc"]==selected_doc_asso_name].index.values[0]
        st.session_state["selected_asso_name"]=selected_doc_asso_name
        
        with st.spinner("Wait for it...", show_time=True):
                for message in process_existing_doc(hash, selected_doc_asso_name, "asso"):
                    st.markdown(message, unsafe_allow_html=True)

        st.session_state["current_asso_in_use"]=selected_doc_asso_name
        st.success(f"✅ Projet asso de la bibliothèque : {selected_doc_asso_name} importé avec succès.")
        with zone_notif_doc_charge_asso:
                st.success(f"""Document chargé: {st.session_state["current_asso_in_use"]}""")
st.divider() 

# --- 2nd step Project proposal template ---
# load an AAP
st.header("Etape 2 - Sélectionnez le formulaire à compléter")
with st.container(key='template_container'):
    col1_aap_loading, col2_aap_loading = st.columns([4, 1], vertical_alignment="center")
    with col1_aap_loading:
        uploaded_aap = st.file_uploader("**Chargez le formulaire à remplir** (format DOCX)",  
                                     accept_multiple_files=True, type=["docx", "json"], key='response_template')
        st.session_state["uploaded_aap"] = uploaded_aap
        if uploaded_aap is not None:
            for uploaded_file in uploaded_aap:
                st.write(f"Fichier {uploaded_file.name} importé avec succès.")
    with col2_aap_loading:
        btn_process_aap=st.button(label="Lancer la réponse", key="btn_response")
    #st.button("Lancer la réponse", key='btn_response')

# Or ask a question about the pp and asso files uploaded or choosen    
    user_query = st.text_input("**Ou posez une question manuellement**", placeholder="Ex : Qui sont les bénéficiaires du projet ?", key='response_text')
    
    col_query1, col_query2, col_query3=st.columns(3, gap="small", vertical_alignment="center", border=False)
    with col_query1:
        btn_process_user_query=st.button(label="Chercher", key="process_user_query",use_container_width=False)
    with col_query2:
        btn_display_sources=st.checkbox(label="Afficher les sources", key="display_sources", )
    with col_query3:
        btn_display_metadata=st.checkbox(label="Afficher les méta données", key="display_metadata",)

# show parameters
    with st.expander("Paramètres"):
        st.slider("Niveau de détail (réponse simple à détaillée)", 1, 5, 3, key='detail_level')

"""     if btn_process_aap:
            if uploaded_aap or response_text:
                st.success("Question traitée avec succès.")
            else:
                st.warning("Veuillez charger un fichier ou saisir une question.")

 """

     # =============afficher les paramètres du rag hybride


# --- Donate and github link buttons ---
col1_link, col2_link = st.columns(2)
with col1_link:
    st.html('<a class="github_link" href="https://github.com/dataforgoodfr/13_ia_financement/" target="_blank">ⓘ À propos de DossierIA+</a>')
with col2_link:
    st.html('<a class="donate_link" href="https://dataforgood.fr/docs/donation/" target="_blank">Faire un don ➚</a> ')


#=================== End UI v2 Streamlit app