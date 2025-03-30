import streamlit as st
from docx import Document
import re
import uuid
from io import BytesIO
import os
from langchain.prompts import PromptTemplate
from langchain.chat_models import ChatOpenAI

# Récupération de la clé API
api_key = os.getenv("OPENAI_API_KEY")

# Vérification de mots liés au format de réponse
def one_of_the_words(text, words_ok, words_ko):
    flag = any(w.lower() in text.lower() for w in words_ok)
    if any(w.lower() in text.lower() for w in words_ko):
        flag = False
    return flag

# Définition de la liste de mots clés pour le foramt de réponse
SIZE_WORDS_OK = ["MAX", "MIN", "CARACT", "CHARACT", "LIGNE", "LINE", "SIGN", "PAGE", "PAS EXC", "NOT EXCEED"]
SIZE_WORDS_KO = ["SIGNAT", "MAXIMI", "MONTH", "MOIS", "ANS", "ANNÉ", "YEAR", "DAY", "JOUR", "DURÉE", "DURATION", "IMPACT", "AMOUNT", "MONTANT"]

# Extraction des questions avec UID dans le document
def extract_questions_with_uid(doc: Document):
    results = {}
    for para in doc.paragraphs:
        if "?" in para.text:
            uid = uuid.uuid4().hex
            question = para.text
            size_info = ''

            if one_of_the_words(para.text, SIZE_WORDS_OK, SIZE_WORDS_KO):
                match = re.search(r'\((.*??)\)', para.text)
                if match:
                    size_info = match.group(1)

            para.add_run(f"\n{uid}")
            results[uid] = [uid, question, size_info, '', '']
    return results, doc

# Page streamlit
def run():
    st.title("Répondre à un Appel à Projet")

    if not api_key:
        st.error("Clé API OpenAI non définie.")
        return

    if "vector_store" not in st.session_state or st.session_state["vector_store"] is None:
        st.warning("Aucun document de projet chargé. Rendez-vous en page 2.")
        return
    else:
        st.info(f"Projet sélectionné : **{st.session_state.get('current_pp_name', 'inconnu')}**")

    uploaded_file = st.file_uploader("📎 Charger un document AAP (.docx)", type=["docx"])

    if uploaded_file:
        doc = Document(uploaded_file)
        questions_dict, modified_doc = extract_questions_with_uid(doc)
        vector_store = st.session_state["vector_store"]
        llm = ChatOpenAI(model="gpt-4o-mini", api_key=api_key)

        for uid, data in questions_dict.items():
            question = data[1]
            docs = vector_store.similarity_search(question)
            context = "\n\n".join(doc.page_content for doc in docs)
            questions_dict[uid][3] = context

            prompt = PromptTemplate.from_template("""
            Tu es chargé de projet pour une association. Tu dois répondre à des appels à projets.
            Utilise uniquement les documents fournis comme contexte.
            Si tu ne trouves pas la réponse, réponds :
            **A REMPLIR PAR L'UTILISATEUR**

            Contexte : {context}
            Question : {question}
            Réponse :
            """)
            response = llm.invoke(prompt.format(question=question, context=context))
            rag_answer = response.content.strip()

            if not rag_answer or len(rag_answer) < 10:
                rag_answer = "**A REMPLIR PAR L'UTILISATEUR**"

            questions_dict[uid][4] = rag_answer

        # Remplacement des UID dans le doc
        for para in modified_doc.paragraphs:
            for uid, data in questions_dict.items():
                if uid in para.text:
                    para.text = para.text.replace(uid, "\n" + data[4])
                if "??" in para.text:
                    para.text = para.text.replace("??", "")

        # Sauvegarde du document modifié 
        output = BytesIO()
        modified_doc.save(output)
        output.seek(0)

        st.success("✅ Réponses générées et insérées dans le document.")

        # Bouton de téléchargement
        st.download_button(
            label="📥 Télécharger le document complété",
            data=output,
            file_name="AAP_repondu.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

         # Aperçu des réponses
        with st.expander("📋 Voir les réponses générées"):
            for uid, data in questions_dict.items():
                st.markdown(f"**Question :** {data[1]}")
                st.markdown(f"**Réponse :** {data[4]}")
                st.markdown("---")
