import os
import re
import streamlit as st
from langchain.document_loaders import PyPDFLoader, Docx2txtLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.prompts import PromptTemplate
from langchain.chat_models import ChatOpenAI
from langchain.schema import Document
from docx import Document as DocxDocument

# Config du template du streamlit 
with st.sidebar:
    st.image("SOS.png")
    st.image("D4G.png")
    st.write("### Projet DataForGood")
st.title("Aide à la rédaction d'AAP avec IA")

# Stockage des fichiers
uploaded_aap = st.file_uploader("Charger le document d'Appel à Projet (PDF/DOCX)", type=["pdf", "docx"])
uploaded_docs = st.file_uploader("Charger les documents de présentation de projets (PDF)", type=["pdf"], accept_multiple_files=True)

# Variables globales
questions = []
answers = []
vector_store = None
api_key = os.getenv("OPENAI_API_KEY")

# Fonction pour sauvegarder temporairement un fichier
def save_uploaded_file(uploaded_file):
    temp_path = os.path.join("temp", uploaded_file.name)
    os.makedirs("temp", exist_ok=True)
    with open(temp_path, "wb") as f:
        f.write(uploaded_file.read())
    return temp_path

# Fonction pour supprimer tous les fichiers temporaires
def clear_uploaded_files():
    if os.path.exists("temp"):
        for file in os.listdir("temp"):
            os.remove(os.path.join("temp", file))
        os.rmdir("temp")
    global questions, answers, vector_store
    questions = []
    answers = []
    vector_store = None
    st.success("Tous les fichiers et données ont été supprimés.")

# Extraction du contenu

def extract_content_text(file_path):
    ext = os.path.splitext(file_path)[-1].lower()
    all_pages_content = []
    
    if ext == ".pdf":
        pdf_loader = PyPDFLoader(file_path)
        documents = pdf_loader.load()
        all_pages_content = [page.page_content for page in documents]
    elif ext == ".docx":
        docx_loader = Docx2txtLoader(file_path)
        documents = docx_loader.load()
        all_pages_content = [page.page_content for page in documents]
    else:
        st.error("Format non pris en charge")
        return []
    
    return all_pages_content


def extract_questions(pages_content):
    pattern = r"<projectQuestion>(.+?)</projectQuestion>"
    return [q.strip() for page_content in pages_content for q in re.findall(pattern, page_content, re.DOTALL)]

if uploaded_aap:
    temp_path = save_uploaded_file(uploaded_aap)
    doc_content = extract_content_text(temp_path)
    questions = extract_questions(doc_content)
    st.success(f"{uploaded_aap.name} chargé avec succès.")
    st.write("**Questions détectées :**")
    for idx, q in enumerate(questions, 1):
        st.write(f"{idx}. {q}")

if uploaded_docs:
    docs = []
    for file in uploaded_docs:
        temp_path = save_uploaded_file(file)
        loader = PyPDFLoader(temp_path)
        docs.extend(loader.load())
    
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200, add_start_index=True)
    all_splits = text_splitter.split_documents(docs)
    
    embeddings = OpenAIEmbeddings(api_key=api_key)
    vector_store = FAISS.from_documents(all_splits, embeddings)
    st.success(f"{len(uploaded_docs)} document(s) de projet chargé(s).")

if st.button("Répondre à l'AAP"):
    if not questions or vector_store is None:
        st.error("Veuillez charger les documents avant de répondre.")
    else:
        template = """Tu es chargé de projet pour une association. Ta mission est de répondre à des appels à projets.
        Tu dois répondre à des {question} en utilisant les documents de contexte qui te sont donnés.
        Si tu ne connais pas la réponse, dis juste que tu ne sais pas, n'essaye pas de répondre absolument.
        Sois clair et précis.
        
        Context : {context}
        Question : {question}
        Réponse utile :"""
        
        prompt = PromptTemplate.from_template(template)
        llm = ChatOpenAI(model="gpt-4o-mini", api_key=api_key)
        answers.clear()
        
        for idx, question in enumerate(questions, 1):
            retrieved_docs = vector_store.similarity_search(question)
            docs_content = "\n\n".join(doc.page_content for doc in retrieved_docs)
            messages = prompt.format(question=question, context=docs_content)
            response = llm.invoke(messages)
            
            answers.append((question, response.content))
            st.write(f"**{idx}. Question :** {question}")
            st.write(f"**Réponse :** {response.content}")

# Exporter en Word
def export_to_word():
    doc = DocxDocument()
    doc.add_heading("Réponses à l'Appel à Projet", level=1)
    
    for idx, (question, answer) in enumerate(answers, 1):
        doc.add_heading(f"Question {idx}", level=2)
        doc.add_paragraph(question)
        doc.add_heading("Réponse", level=2)
        doc.add_paragraph(answer)
    
    temp_file = "reponses_aap.docx"
    doc.save(temp_file)
    
    with open(temp_file, "rb") as f:
        st.download_button(
            label="Télécharger les réponses en Word",
            data=f,
            file_name="reponses_aap.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if answers:
    export_to_word()

if st.button("Supprimer tous les fichiers et données"):
    clear_uploaded_files()