import streamlit as st
from rag_pipelines import process_new_doc, process_existing_doc, QA_pipeline, update_hybrid_rag_wrapper, adjust_resp
from graphrag_retriever import load_knowledgeGraph_vis
from io import StringIO
import asyncio
import nest_asyncio
import os
import pandas as pd
import json
from read_answer_aap import Read_Questions_in_docx, Write_Answers_in_docx
from pathlib import Path
import traceback

# pour télécharger l'AAP en docx
from docx import Document
from docx.shared import Pt
from io import BytesIO

# Patch torch.classes pour éviter l'erreur de Streamlit
import torch



#=========corrections incompatibilités streamlit / event loop
#=============== nécessaire pour streaming des réponses graphRAG
# Patch the event loop to allow nested async calls
nest_asyncio.apply()
#==========fin

#=================== Correction conflit event loops torch/streamlit
# Save the original __getattr__ method
original_getattr = torch._classes._Classes.__getattr__

# Define a patched version to handle __path__
def patched_getattr(self, attr):
    if attr == "__path__":
        # Explicitly block access to __path__
        raise AttributeError(f"'{self.__class__.__name__}' has no attribute '__path__'")
    return original_getattr(self, attr)

# Apply the patch
torch._classes._Classes.__getattr__ = patched_getattr
#===============fin



# Get the directory of the current script (e.g., app.py)
SCRIPT_DIR = Path(__file__).parent.resolve()



def stream_pathRAG_response(stream_resp, response_container):
    async def stream_response():
        response_buffer = StringIO()        
        
        # Get the existing event loop
        loop = asyncio.get_event_loop()
        
        # Process the async generator
        async for chunk in stream_resp["response_stream"]:
            response_buffer.write(chunk)
            response_container.markdown(response_buffer.getvalue())

        st.session_state["full_response"] = response_buffer.getvalue()
        response_buffer.close()


    # Run in Streamlit's existing event loop
    loop = asyncio.get_event_loop()
    loop.run_until_complete(stream_response())

def stream_hybridRAG_response(stream_resp, response_container):
    # 1. Créer un buffer pour accumuler la réponse
    response_buffer = StringIO()
    response_container = st.empty()  # Conteneur vide pour mise à jour dynamique

    for chunk in stream_resp["response_stream"]:
    # Ajouter le token au buffer
        response_buffer.write(chunk)
        # Mettre à jour le conteneur avec le contenu accumulé
        response_container.markdown(response_buffer.getvalue())

    # 3. Récupérer la réponse complète
    st.session_state["full_response"] = response_buffer.getvalue()
    response_buffer.close()


# chargement du log des documents dispo en DB
def load_pp_traces(doc_category: str):
    st.markdown(f"### Documents {doc_category.upper()} disponibles:", unsafe_allow_html=True)


    # charger les hashes
    files_paths=[SCRIPT_DIR/ "hybridrag_hashes.json", SCRIPT_DIR/"graphrag_hashes.json"]
    table=pd.DataFrame([])

    for file_path in files_paths:
        exising_hashes={}
        if os.path.exists(file_path):                
            t=pd.read_json(file_path,).T
            table=pd.concat([table, t], axis=0)

    # préparer les hashes
    if len(table)>0:
        # enlever les lignes en double hybrid rag/ graph rag
        table=table.drop_duplicates(subset=["Nom du doc","Titre auto","Taille du texte (en car)"])
        table=table.sort_values("Date de création", ascending=False)
        table=table[table["doc_category"]==doc_category].drop(columns=["rag_type"])
        table['Date de création'] = pd.to_datetime(table['Date de création']).dt.floor('s')

        st.dataframe(table.reset_index(drop=True))
        st.markdown("---")
        return table
        
    else:
        st.write("Aucun document pour le moment")
        st.markdown("---")
        return []

def generate_docx(data):
    doc = Document()
    
    # Style du titre
    title = doc.add_heading('Questions/Réponses', 0)
    title.style.font.size = Pt(14)
    
    for item in data:
        # Ajout de la question
        doc.add_heading(item['question'], level=1)
        
        # Ajout de la réponse
        doc.add_paragraph(item['response'])
        
        # Ligne de séparation
        doc.add_paragraph().add_run().add_break()
    
    # Sauvegarde en mémoire
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


def main():
    
    def build_ui_pp_asso(doc_category: str):
        
        # charger traces des PP existantes:
        df_existing_docs= load_pp_traces(doc_category)



        #========zone document en cours d'usage
        zone_notif_doc_charge=st.empty()
        if f"current_{doc_category}_in_use" in st.session_state:
            zone_notif_doc_charge.success(f"""Document chargé: {st.session_state[f"current_{doc_category}_in_use"]}""")
        else:
            zone_notif_doc_charge.warning(f"Veuillez charger un document {doc_category.upper()}")

        #======== UI Charger un nouveau doc    
        # UI définitive        
        st.write(f'### Charger un nouveau document {doc_category.upper()}')
        uploaded_doc= st.file_uploader(label=f"Zone de chargement {doc_category.upper()}", type=["pdf","docx"], accept_multiple_files=True, key="uploader_doc")
        st.session_state[f"uploaded_{doc_category}"]=uploaded_doc

        session_state=st.session_state
        print(session_state)
        col1_new_doc, col2_new_doc = st.columns(2, vertical_alignment="bottom")
        with col1_new_doc:
            doc_named_user=st.text_input(label=f"Nom du document {doc_category.upper()}", placeholder="Saisie obligatoire")
        with col2_new_doc:
            btn_new_doc= st.button("Traiter", key=f'btn_{doc_category}')
            st.session_state["btn_new_doc"]=btn_new_doc


        st.markdown('<hr>', unsafe_allow_html=True)


        #===============================================





        #========UI Utiliser un doc existant
        st.write(f'### Utiliser un document {doc_category.upper()} disponible')
        col4_existing_doc, col5_existing_doc = st.columns(2, vertical_alignment="bottom", )
        with col4_existing_doc:
            list_doc_names=[]
            if len(df_existing_docs)>0:
                list_doc_names=df_existing_docs[f"Nom du doc"].unique()
            
            selected_doc_name= st.selectbox(label=f"Choisir un document {doc_category}", options=list_doc_names)

            # col1_load_doc,col2_load_doc=st.columns(2, gap="small")
            # with col2_load_doc:
            btn_load_existing_doc=st.button(label="Charger", key='load_existing_doc')
                
        # with col5_existing_doc:
        #     btn_existing_doc= st.button(f"Charger", key=f"btn_process_{doc_category}")
        #================================================




        #========Interactions nouveau doc
        # définitif
        if uploaded_doc is not None and doc_named_user!="" and btn_new_doc:
            st.session_state[f"uploaded_{doc_category}"]=uploaded_doc
            with st.spinner("Wait for it...", show_time=True):
                
                for message in process_new_doc(uploaded_doc, doc_named_user, doc_category):
                    st.markdown(message, unsafe_allow_html=True)

            st.session_state[f"current_{doc_category}_in_use"]=doc_named_user
            st.success(f"✅ Traitement de {doc_named_user} terminé")
            with zone_notif_doc_charge:
                st.success(f"""Document chargé: {st.session_state[f"current_{doc_category}_in_use"]}""")

        #===============================================
        
        
        #========Interactions PP existant
        
        elif btn_load_existing_doc and selected_doc_name:
            hash=df_existing_docs[df_existing_docs[f"Nom du doc"]==selected_doc_name].index.values[0]
            st.session_state["selected_pp_name"]=selected_doc_name
            
            with st.spinner("Wait for it...", show_time=True):
                for message in process_existing_doc(hash, selected_doc_name, doc_category):
                    st.markdown(message, unsafe_allow_html=True)
            
            st.session_state[f"current_{doc_category}_in_use"]=selected_doc_name
            st.success(f"✅ Chargement de {selected_doc_name} terminé")
            with zone_notif_doc_charge:
                st.success(f"""Document chargé: {st.session_state[f"current_{doc_category}_in_use"]}""")


        #===============================================


    st.set_page_config(
        page_title="IA pour répondre aux appels à projets",
        layout="wide"
    )
    
        # --- CSS personnalisé avec background bleu ciel clair ---
    st.markdown(
        """
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@400;700&family=Roboto:wght@400;500&display=swap');
        
        /* Fond général et header */
        html, body, [data-testid="stAppViewContainer"] {
            font-family: 'Roboto', sans-serif;
            
            



        }
        
        
        /* Bannière principale avec image centrée */
        .main-title {
            text-align: center;
            padding: 15px;
            border-radius: 5px;
            margin-bottom: 20px;
            color: #003366;
        }
        .main-title h1 {
            font-family: 'Montserrat', sans-serif;
            font-size: 3rem;
            margin: 10px 0 0;
        }
        .main-title p {
            font-size: 1.2rem;
            margin: 8px 0 0;
        }
        /* Style pour l'image de la bannière : centrée et occupant 80% de la largeur */
        .main-title img {
            display: block;
            margin-left: auto;
            margin-right: auto;
            width: 80%;
            max-height: 300px;
            object-fit: contain;
        }
        
        /* Sidebar : on applique le dégradé à l'ensemble de la sidebar 
        et on centre les images */
        [data-testid="stSidebar"] {
        /*background: linear-gradient(90deg, #8EC0FA 30%, #FFDDE0 90%) !important;*/
        background: linear-gradient(90deg, #0D5C78 0%, #8EC0FA 50%, #2A6351 100%);
        padding: 20px;
        border-radius: 0 0 12px 0;
        }
        [data-testid="stSidebar"] img {
        display: block;
        margin-left: auto;
        margin-right: auto;
        }
        
        /* Bouton personnalisé */
        .stButton > button {
            linear-gradient(90deg, #6495ED 30%, #FFB6C1 90%) !important;
            color: #003366;
            border-radius: 8px;
            border: 2px solid #CCCCFF;
            padding: 10px 24px;
            font-size: 1.1rem;
            font-weight: 600;
            transition: background-color 0.3s ease, transform 0.2s ease;
        }
        .stButton > button:hover {
            background: linear-gradient(90deg, #6495ED 30%, #FFB6C1 90%) !important;
            transform: scale(1.03);
        }
        
        /* Espacement des colonnes */
        div[data-testid="column"] {
            padding: 10px;
        }
        </style>
        """,
        unsafe_allow_html=True
    )


    # Bannière principale affichant le titre et un sous-titre
    st.markdown(
        """
        <div class="main-title">
            <h1>IA pour répondre aux appels à projets</h1>
            <p>Optimisez vos réponses grâce à l'intelligence artificielle</p>
        </div>
        """,
        unsafe_allow_html=True
    )


    # menu
    st.sidebar.title("Sommaire")
    pages_map={
        "Guide d'utilisation": 0, 
        "Chargement PP": 1,
        "Chargement asso": 2,
        "Remplir un AAP": 3, 
        "Paramètres": 4, 
        "Graph de connaissances": 5
    }
    pages=[k for k in pages_map.keys()]

    #persistance
    index_sommaire=0
    if 'sidebar' in st.session_state:
        index_sommaire= pages_map[st.session_state["sidebar"]]
    
    page=st.sidebar.radio("Aller vers", pages, key="sidebar", index=index_sommaire)
    
    
    with st.sidebar:
        try:
            img_1=SCRIPT_DIR/ "SOS.png"
            img_2=SCRIPT_DIR/ "D4G.png"
            st.image(img_1)
            st.image(img_2)
        except Exception as e:
            st.write(e)
            st.write(img_1)
            st.write(img_2)

        st.write("### Projet DataForGood")
        

    # page d'accueil
    if page == pages[0]:
        import page1_intro
        page1_intro.run()


    # page chargement de la PP
    if page == pages[1]:
      
        build_ui_pp_asso('pp')


    if page == pages[2]:
        build_ui_pp_asso('asso')


    # =============== Q/A AAP ou directe
    if page == pages[3]:
        st.write("#### Charger un AAP")
        with st.expander("Zone formulaire AAP"):                    
            uploaded_aap = st.file_uploader(
                label="Glisser ou charger un formulaire", 
                type=["docx", "json"], 
                accept_multiple_files=False, 
                key="uploader_aap"
            )
            st.session_state["uploaded_aap"] = uploaded_aap

            btn_process_aap=st.button(label="Traiter", key="process_aap")

        st.markdown("------------", unsafe_allow_html=True)

        st.write("#### Saisie manuelle")
        with st.expander("Zone saisie de question"):
            user_query = st.text_input(label="Votre question", placeholder="")
            user_query_size = st.text_input(label="Taille de réponse souhaitée", placeholder="")

            col_query1, col_query2, col_query3=st.columns(3, gap="small", vertical_alignment="center", border=False)

            with col_query1:
                btn_process_user_query=st.button(label="Chercher", key="process_user_query",use_container_width=False)                
            with col_query2:
                btn_display_sources=st.checkbox(label="Afficher les sources", key="display_sources", )
            with col_query3:
                btn_display_metadata=st.checkbox(label="Afficher les méta données", key="display_metadata",)

        st.markdown("------------", unsafe_allow_html=True)

        # =============afficher les paramètres du rag hybride 
        st.markdown("#### Paramètres du RAG hybride:")
        # with st.expander("Zone formulaire AAP"):
        # buf1, buf2,  col_reranker_select, col_top_k=st.columns([1, 1, 1, 1])
        def reranker_template(reranker_holder):
            # '#fa8072e6': salmon, "#008000d9": green
            bkg_color= '#fa8072e6' if reranker_holder=='specialized' else "#008000d9"
            return f"""<span 
                style='
                    background-color: {bkg_color};
                    border-radius: 4px;
                    padding: 2px 5px 2px 5px;
                    color: white;'
                > {reranker_holder}
            </span>"""

        if "selected_reranker" in st.session_state and st.session_state["selected_reranker"]!="":
            # with col_reranker_select:
            st.markdown(f"""* Reranker utilisé: **{reranker_template(st.session_state["selected_reranker"])}**""", unsafe_allow_html=True)
        else:
            # with col_reranker_select:
            st.markdown(f"""* Reranker utilisé: {reranker_template("specialized")}""",
             unsafe_allow_html=True)

        if "top_k_docs_user" in st.session_state:
            # with col_top_k:
            st.markdown(f"* Top_k documents utilisé: **{st.session_state['top_k_docs_user']}**", unsafe_allow_html=True)
        else:
            # with col_top_k:
            st.markdown(f"* Top_k documents utilisé: **{10}**", unsafe_allow_html=True)

        st.markdown("----", unsafe_allow_html=True)


        queries = []

        # === Saisie manuelle ===
        if btn_process_user_query and user_query.strip() != "":
            queries = [{"question": user_query, "size_answer": user_query_size}]
            st.session_state["trigger_query"] = False  # reset

        # === Traitement AAP (json/docx) ===
        elif btn_process_aap and uploaded_aap is not None:
            try:
                st.session_state["trigger_aap"] = False  # reset

                if uploaded_aap.name.endswith(".docx"):
                    st.info("📄 Traitement automatique du fichier AAP...")


                    # Construction du chemin absolu pour AAP, LOG et output_aap
                    source_aap = SCRIPT_DIR / "AAP/" # répertoire pour l'AAP source (non rempli)
                    hidden_log = SCRIPT_DIR / "LOG/" # répertoire pour l'AAP avec UID + le log file (pas accessible par l'utilisateur)
                    output_aap = SCRIPT_DIR / "output_aap/" # répertoire où sera placé l'AAP avec réponses

                    # Nettoyage du nom de fichier
                    safe_name = Path(uploaded_aap.name).name  # Exemple : "PU_P01_AAP01 - sample.docx"

                    # Construction du chemin complet avec chemin absolu
                    file_path_in = source_aap / safe_name 

                    # Création du dossier parent si nécessaire
                    Path(source_aap).mkdir(parents=True, exist_ok=True)
                    Path(hidden_log).mkdir(parents=True, exist_ok=True)
                    Path(output_aap).mkdir(parents=True, exist_ok=True)

                    print(f"Chemin absolu avec fichier : {file_path_in}")
                    print(f"Chemin absolu sans fichier: {hidden_log}")
                    print(f"Chemin absolu sans fichier : {output_aap}")

                    # Suppression des anciens fichiers source (de source_aap) au cas où pas été supprimés dans traitements anciens
                    for old_file_path in source_aap.glob('*.*'):
                        try:
                            old_file_path.unlink()
                            print(f"Suppression fichiers AAP vierges préalable avant fonction Read : {old_file_path}")
                        except Exception as e:
                            print(f"Erreur lors de la suppression de l'ancien fichier source {old_file_path} : {e}")

                    # Suppression des anciens fichiers "avec UID" (de hidden_log) au cas où pas été supprimés dans traitements anciens
                    for old_file_path in hidden_log.glob('*.*'):
                        try:
                            old_file_path.unlink()
                            print(f"Suppression fichiers avec UID préalable avant fonction Read : {old_file_path}")
                        except Exception as e:
                            print(f"Erreur lors de la suppression de l'ancien fichier avec UID {old_file_path} : {e}")

                    # Suppression des anciens fichiers AAP avec réponses (de output_aap) au cas où pas été supprimés dans traitements anciens
                    for old_file_path in output_aap.glob('*.*'):
                        try:
                            old_file_path.unlink()
                            print(f"Suppression fichiers AAP avec réponse et fichier QA préalable avant fonction Read : {old_file_path}")
                        except Exception as e:
                            print(f"Erreur lors de la suppression de l'ancien fichier AAP avec réponses {old_file_path} : {e}")


                    #==========================                
                
                    with open(file_path_in, "wb") as f:
                        f.write(uploaded_aap.getbuffer())

                    with st.spinner("🔍 Extraction des questions en cours..."):
                        extracted_questions = Read_Questions_in_docx(
                            PathFolderSource=source_aap,
                            PathForOutputsAndLogs=hidden_log,
                        )

                    st.success("✅ Extraction terminée")
                    st.write(f"Nombre de questions détectées : {len(extracted_questions)}")
                    queries = extracted_questions
            except Exception as e:
                print(e)
                print(traceback.format_exc())

        #### rappel de la dernière Q/A traitée
        response_container = st.empty() 
        if "full_response" in st.session_state:         
            query_value=""
            if "user_query" in st.session_state:
                query_value=st.session_state["user_query"]
                                              
            QA_pair=f"""<h4>Question:</h4>
                {query_value} 
                <br><br>
                <h4>Réponse:</h4>
            {st.session_state["full_response"]}
            """
            response_container.markdown(QA_pair, unsafe_allow_html=True)


        #============Gestion des interactions===========
        # vérifier qu'une question est posée
        start_process=False
        # cas AAP chargé & préparé
        if len(queries) and btn_process_aap==True>0:
            start_process=True
        
        # cas question posée manuellement
        elif (user_query!="" and btn_process_user_query):
            start_process=True
            if "all_responses_to_write" in st.session_state:                
                del st.session_state["all_responses_to_write"]
        
        if start_process:
            #= reset de la réponse affichée
            response_container.markdown("")
            st.session_state["user_query"]=user_query
            #====== déterminer si requête manuelle ou process AAP
            #1. requête manuelle
            if btn_process_user_query:
                queries=[{"question": user_query, "size_answer": user_query_size}] 



            #======parcourir les questions et les transmettre à QA pipeline
            all_responses_to_write = []
            with st.spinner("Wait for it...", show_time=True):
                for resp in QA_pipeline(queries):
                    
                    # QA_pipeline va retourner plusieurs types de messages

                    # 1. la réponse sous forme de flux à diffuser sur l'UI, + la question initiale posée
                    if isinstance(resp, dict) and "response_stream" in resp:
                        # rappel de la question
                        st.markdown(f"#### Question:\n", unsafe_allow_html=True)
                        st.markdown(resp["question"])
                        st.session_state["user_query"]=resp["question"]

                        st.markdown(f"#### Réponse:\n", unsafe_allow_html=True)                
                        response_container = st.empty()

                        # diffusion PathRag
                        if "pathrag_stream" in resp:
                            stream_pathRAG_response(stream_resp=resp, response_container = response_container)
                        
                        # diffusion hybrid RAG
                        elif "hybridrag_stream" in resp:       
                            stream_hybridRAG_response(stream_resp=resp, response_container = response_container)

                        
                        st.markdown("-----", unsafe_allow_html=True)

                        
                        
                    # 2. les sources utilisées par le rag (seulement pour requete manuelle)
                    elif isinstance(resp, dict) and 'sources' in resp and btn_process_user_query and btn_display_sources:
                        st.markdown(f"#### Sources:\n", unsafe_allow_html=True)
                        
                        for source in resp["sources"]:
                            st.markdown(f"**Source**:<br>{source[0]}", unsafe_allow_html=True)
                            st.markdown(f"**Score**: {source[1]}", unsafe_allow_html=True)

                    # 3. les métadonnées (uid, question, type), et la réponse complète du flux 1 ci dessus
                    elif isinstance(resp, dict) and 'uid' in resp:
                        resp["response"]=st.session_state["full_response"]
                        size_answer = resp["size_answer"] 
                        if size_answer!="": 
                            adjusted_resp=adjust_resp(resp["response"], size_answer)  
                            resp["adjusted_resp"]=adjusted_resp 

                            st.markdown(f"#### Réponse ajustée:\n", unsafe_allow_html=True)             
                            response_container = st.empty()  
                            st.markdown(resp["adjusted_resp"], unsafe_allow_html=True)  
                        else:
                            resp["adjusted_resp"]=""


                        if  btn_process_user_query and btn_display_metadata:
                            st.markdown(f"**Metadata**:", unsafe_allow_html=True)
                            st.json(resp)

                        all_responses_to_write.append(resp) 

                        
                    # 4. diffusion d'un feedback str
                    elif isinstance(resp, str):                    
                        st.markdown(f"{resp}", unsafe_allow_html=True)


            # ✅ Remplir AAP: Une seule écriture à la fin
            if (btn_process_aap and all_responses_to_write):
                st.session_state["all_responses_to_write"]=all_responses_to_write

                output_file_path, qa_file_path, aap_filename, qa_filename = Write_Answers_in_docx(
                    PathFolderSource=hidden_log,
                    PathForOutputsAndLogs=output_aap,
                    List_UIDQuestionsSizeAnswer=all_responses_to_write
                )

                st.session_state["output_file_path"]=output_file_path
                st.session_state["qa_file_path"]=qa_file_path
                st.session_state["aap_filename"]=aap_filename
                st.session_state["qa_filename"]=qa_filename


                with output_file_path.open (mode="rb+") as file:
                    aap_data = file.read()
                    
                btn_down_aap = st.download_button(
                    label="⬇️ Télécharger le document avec réponses",
                    data=aap_data,
                    file_name=aap_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    # key="download_aap_btn",  # Clé unique pour le widget
                    help="Cliquez pour télécharger le document Word avec les réponses intégrées"
                )
                
                if btn_down_aap:
                    st.success("Téléchargement lancé avec succès!")



                # Bouton pour télécharger le fichier Q&A
                with qa_file_path.open (mode="rb+") as file:
                    qa_data=file.read()

                btn_down_qa= st.download_button(
                    label="⬇️ Télécharger le fichier Q&A",
                    data=qa_data,
                    file_name=qa_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",

                )
                if btn_down_qa:
                    st.success("Téléchargement lancé avec succès!")
            
                st.success("Document prêt pour téléchargement !")                

        # maintenir la consistance de la page (Q/R et boutons téléch AAP)
        if "all_responses_to_write" in st.session_state and btn_process_aap==False:

            output_file_path=st.session_state["output_file_path"]
            qa_file_path=st.session_state["qa_file_path"]
            aap_filename=st.session_state["aap_filename"]
            qa_filename=st.session_state["qa_filename"]


            with output_file_path.open (mode="rb+") as file: 
                aap_data = file.read()
                
            btn_down_aap = st.download_button(
                label="⬇️ Télécharger le document avec réponses",
                data=aap_data,
                file_name=aap_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                # key="btn_down_aap",  # Clé unique pour le widget
                help="Cliquez pour télécharger le document Word avec les réponses intégrées"
            )
            
            if btn_down_aap:
                st.success("Téléchargement lancé avec succès!")


            # Bouton pour télécharger le fichier Q&A
            with qa_file_path.open (mode="rb+") as file: 
                qa_data=file.read()

            btn_down_qa= st.download_button(
                label="⬇️ Télécharger le fichier Q&A",
                data=qa_data,
                file_name=qa_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",

            )
            if btn_down_qa:
                st.success("Téléchargement lancé avec succès!")
                
            



    # ===============Paramètres
    elif page == pages[4]:
        
        st.markdown("### Paramètres de la chaîne hybride", unsafe_allow_html=True)
        st.markdown("""
            #### Reranker:
            Sert à noter la pertinence des fragments de documents correspondants à la question, seuls les plus pertinents sont envoyés au LLM.
                    
            * Le reranker 'specialized' est rapide et gratuit.
            * Le reranker 'llm' (gpt4o-mini adapté) est lent et payant, mais plus précis.
                    
            Changer de reranker si la réponse produite n'est pas satisfaisante

            #### Top K documents:
            Sert à déterminer le nombre de fragments de documents sélectionnés pour constituer le contexte présenté au LLM, après notation et classement.
                    
            * Réduire ce nombre permet d'augmenter la capacité du LLM à se concentrer sur un contexte précis, mais réduit aussi la diversité des sources, privant le LLM de la bonne information.
            * Augmenter ce nombre permet d'augmenter la diversé des sources et informations nécessaires pour répondre à la réponse, mais peut introduire d'avantage 
            de confusion pour le LLM qui se voit présenté un contexte plus large.

            Changer ce nombre si la réponse n'est pas satisfaisante        
            
            **Conseil**: Changer de reranker en priorité si la réponse produite n'est pas satisfaisante, et éviter de dépasser top K 14
                                        
        """)
                

        # rappel du reranker
        last_reranker=""
        reranker_map={"specialized":0, 'llm':1}
        if "selected_reranker" in st.session_state and st.session_state["selected_reranker"]!="":
            last_reranker=st.session_state["selected_reranker"]
        
        
        st.markdown("---------")
        # st.write(f"Dernier reranker utilisé: {last_reranker}")
        @st.dialog("Validez votre choix")
        def vote():
            st.markdown(f"""Cliquer sur <span 
                style="    
                    border-radius: 6px;
                    border: 2px solid;
                    padding: 6px 8px 6px 8px;
                    margin: 0px 4px 0 4px;
                    border-color: lightgray;
                    "
            >Mise à jour</span> pour activer le changement""", 
            unsafe_allow_html=True)

        reranker=st.selectbox(
            "Changer de reranker",
            options=["specialized", "llm"],
            on_change=vote
        )


        # rappel du top K
        top_k_user_input=st.empty()
        if 'top_k_docs_user' in st.session_state:                        
            # "avant", st.session_state["top_k_docs_user"]
            top_k_docs_selected= top_k_user_input.number_input(label="Top K documents", key="top_k_user_input", value=st.session_state["top_k_docs_user"], min_value=4, max_value=18)
            
        else:
            
            top_k_docs_selected= top_k_user_input.number_input(label="Top K documents", key="top_k_user_input",  value=10, min_value=4, max_value=18)
            
            
        st.session_state["top_k_docs_user"]=top_k_docs_selected
        # "après", st.session_state["top_k_docs_user"]

        change_hybridrag_params=st.button("Mise à jour", key="update_hybrid_rag_params")

        if change_hybridrag_params:
            st.session_state["selected_reranker"]= reranker
            feedback= update_hybrid_rag_wrapper(reranker, top_k=top_k_docs_selected)
            st.write(f"""{feedback}""")
   
    # ==============Graph
    elif page == pages[5]:

        # cas où aucun PP n'est chargé
        

        for feedback in load_knowledgeGraph_vis():
            if isinstance(feedback, str):
                st.markdown(feedback)
            elif isinstance(feedback, tuple):
                graphvis_path, graphvis_name=feedback
                
                st.markdown(f"Document: {graphvis_name}")
                with open(graphvis_path, "r", encoding="utf-8") as file:
                    html_content = file.read()
                st.components.v1.html(html_content, height=1200, scrolling=True)




if __name__ == "__main__":
    main()