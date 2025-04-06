
#@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@ TO DO IN READ @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@
# Bug1 de AAP PU_P01_AAP05-with UID.docx : comme il y a une colonne à gauche, il y a une cellule vide dessous 
# mais il met l'UID dans la cellule de la question car il considère être en type1 et + avec aucune colonne vide à droite
# Il faudrait lui faire mettre dans la cellule du dessous. POur éviter ça, il faudrait identifier si il y a une cellule fusionnée
# à gauche et en ce cas, s'il n'y a pas de cellule à droite, voir si cellule vide dessous et dépôt UID
# Bug2 de AAP PU_P01_AAP05-with UID.docx : Dans une cellule fusionnée où il y a une question, il considère qu'il y a deux fois la
# question et il ajoute 2 fois l'UID. POur éviter cela, il faudrait vérifier que le contenu de la cellule suivante ayant une question
# n'est pas égal au contenu de la cellule précédente => If contenu préc != contenu acutelle , ajout question et metre uid
# Bug AAP "PU_P01_AAP07 Section 1-4_with??" 2°) des tableaaux Type 2 avec ? non vus = tableau du Q1.6 
# Mieux voir les (500 words) des tableaux = ceux dans la cellule en dessous ou dans la cellule à côté
# Mieux voir les (500 words) des non-tableaux = ceux dans paragrphes en dessous
# A LA FIN DIFFUSER A MARINE EN ENLEVANT LES TO DO
#@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@ END TO DO IN READ @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@

#@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@ TO DO IN WRITE @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@
# if key not found, log an error
# PRIO 01 : modifier code pour prendre en compte la liste de dictionnaires de Aghiles
# {
#   'uid': 'uid',
#   "question": ...
#   "enhanced_question": ...,
#   "question_is_open": close/open
#   "question_on_asso": yes/no
# }
#@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@ END TO DO IN READ @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@

# #@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@ END TESTS & IMPROVEMENTS NEEDED @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@

#@@@@@@@@@@@@@@@@@@@@@@@@@@@@ "SCAN PARAGRAPHS & TABLES TOP DOWN IN A DOCX DOCUMENT" FUNCTION @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@
def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document
    order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    import docx
    from docx.document import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph

    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

#@@@@@@@@@@@@@@@@@@@@@@@@@@@@ END OF "SCAN PARAGRAPHS & TABLES TOP DOWN IN A DOCX DOCUMENT" FUNCTION @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@



#@@@@@@@@@@@@@@@@@@@@@@@@@@@@ "ONE OF THE WORDS IS IN THE PARAGRAPH" FUNCTION @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@
def OneOfTheWords_Is_InTheParagraph (TheText, list_of_Words_OK, list_of_Words_KO):
    """
    This function verifies if one of the words of a list of words is in a paragraph
        
    Args:
        docpara : the paragraph in which we verify
        list_of_Words_OK : List of words that we want to check if they are present in the paragraph
        list_of_Words_KO : List of words that indicate wrong interpretation of the Words OK
            in other terms, if we find a word of list_of_Words_OK in the paragraph but we also 
            find there a word of list_of_Words_KO, it disqualifies the 1st finding and we consider no presence of the word in the paragraph
            e.g. : we want to find the word meaning "maximum" so we look for "MAX" (OK list) because maximum is often written "max."
            we find it, but find also "MAXIMIZE" (KO list), in this case MAX does not means "MAXIMUM" but it is part of "MAXIMIZE" which is wrong for our quest
            
    Returns:
        The function returns True if a word from list_of_Words_OK is found 
        and no word from list_of_Words_KO is found
         Else it returns False
    """
    FlagWord_OK = False # by default, we consider no word found in the paragraph
        #============== 1 - TREATMENT OF "LIST_OF_WORDS_OK" ================================================    
    for Theword in list_of_Words_OK: # We look for words from the list list_of_Words_OK
        #if the word in lowercase is in the text in lowercase, we have found one matching word
        if re.search(Theword.lower(), TheText.lower(), flags=0)!= None:
            FlagWord_OK = True

        #============== 2 - TREATMENT OF "LIST_OF_WORDS_KO" ================================================    
    for Theword in list_of_Words_KO: # Now we look for words from the list list_of_Words_KO
        #if the word in lowercase is in the text in lowercase, we have found one matching word
        if re.search(Theword.lower(), TheText.lower(), flags=0)!= None:
            FlagWord_OK = False # the Word of list_of_Words_KO disqualifies the word of list_of_Words_OK
            #if the keyword in lowercase is in the text in lowercase, we have found one matching word
    return FlagWord_OK # return True if found or False if not found

#@@@@@@@@@@@@@@@@@@@@@@@@@@@@ END OF "ONE OF THE WORDS IS IN THE PARAGRAPH" FUNCTION @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@



#@@@@@@@@@@@@@@@@@@@@@@@@@@@@ "INSERT TEXT IN ONE PARAGRAPH IN FULL TEXT (NO TABLE)" FUNCTION @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@
def Insert_Text_Paragraph (block_item, TextStart, TextEnd):
    """
    This function inserts text into a paragraph of Word docx at the beginning and at the end of the paragraph
    This function works for paragraphs in full text (i.e. not inside tables)
    To insert text inside cells of tables, another code is required
    This code allows to insert the text in the paragraph using "replace" function which is the only way to do it
    without loosing the initial look & feel of the texte (size, font, color,..)
    because any other way of changing the text of a paragraph in docx Word will unfortunately loose all of that

    Args:
        block_item : the paragraph in which we insert the text
        TextStart : the text to be inserted at the beginning of the paragraph
        TextEnd : the text to be inserted at the end of the paragraph
            
    Returns:
        The function returns nothing
        but modifies the paragraph by adding text
    """
    if block_item.runs == []: # if the paragraph has no run
        block_item.text = TextStart + block_item.text + TextEnd # we manage at text level
    else: # if the paragraph has at least 1 run, we manage at run level
        # insert the start text
        block_item.runs[0].text = block_item.runs[0].text.replace("", TextStart,1) 
        # then insert the end text
        NbRuns = block_item.runs.__len__()
        block_item.runs[NbRuns-1].text = block_item.runs[NbRuns-1].text.replace(block_item.runs[NbRuns-1].text, block_item.runs[NbRuns-1].text + TextEnd,1)

    return
#@@@@@@@@@@@@@@@@@@@@@@@@@@@@ END OF "INSERT TEXT IN ONE PARAGRAPH IN FULL TEXT (NO TABLE)" FUNCTION @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@



#@@@@@@@@@@@@@@@@@@@@@@@@@@@@ "DELETE TEXT IN ONE PARAGRAPH IN FULL TEXT (NO TABLE)" FUNCTION @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@
def Delete_Text_Paragraph (block_item, Text_to_delete):
    """
    This function deletes text in a paragraph of Word docx
    This function works for paragraphs in full text (i.e. not inside tables)
    To delete text inside cells of tables, another code is required
    This code allows to delete the text in the paragraph using "replace" function which is the only way to do it
    without loosing the initial look & feel of the texte (size, font, color,..)
    because any other way of changing the text of a paragraph in docx Word will unfortunately loose all of that

    Args:
        block_item : the paragraph in which we insert the tags
        Text_to_delete : the text to be deleted in the paragraph
            
    Returns:
        The function returns nothing
        but modifies the paragraph by deleting text
    """
    Text_to_delete2 =""
    if block_item.runs == []: # if the paragraph has no run
        block_item.text = block_item.text.replace(Text_to_delete, "") # suppress the Text_to_delete but loose the initial look & feel (format) of the paragraph)

    else: # if the paragraph has at least 1 run, we manage at run level
        # we have to create Text_to_delete2 because re.search will not work with simple ? or > or < or /
        if Text_to_delete =="??": 
            Text_to_delete2 = r'\?\?'
        if Text_to_delete =="<>":
            Text_to_delete2 = r'\<\>'
        if Text_to_delete =="</>":
            Text_to_delete2 = r'\<\/\>'
        
        NbRuns = block_item.runs.__len__()
        for i in range(NbRuns):  # Loop through all runs in the paragraph
            MyRun = block_item.runs[i]
            if Text_to_delete2 !="": # if it is "??" or "<>"" or "</>", we use Text_to_delete2
                if re.search(Text_to_delete2, MyRun.text, flags=0)!= None :
                    MyRun.text = MyRun.text.replace(Text_to_delete, '',1) 
            else: # if it is NOT "??" or "<>"" or "</>", use Text_to_delete
                if re.search(Text_to_delete, MyRun.text, flags=0)!= None :
                    MyRun.text = MyRun.text.replace(Text_to_delete, '',1) 
        if Text_to_delete2 !="" and re.search(Text_to_delete2, block_item.text, flags=0)!= None :# if after run treatment, the text to delete not deleted (runs cut le texte to delete in 2)
            # Try to save the former format of paragraph by saving the format of the last run
            NbRuns = block_item.runs.__len__()
            MyFontName = block_item.runs[NbRuns-1].font.name
            MyFontSize = block_item.runs[NbRuns-1].font.size
            MyFontBold = block_item.runs[NbRuns-1].font.bold
            MyFontItalic = block_item.runs[NbRuns-1].font.italic
            MyFontUnderline = block_item.runs[NbRuns-1].font.underline
            MyFontColor = block_item.runs[NbRuns-1].font.color.rgb
            block_item.text = block_item.text.replace(Text_to_delete, "") # suppress the Text_to_delete but loose the initial look & feel (format) of the paragraph)
            # Try to re establish the former style
            NbRuns = block_item.runs.__len__()
            for i in range(NbRuns):  # Loop through all runs in the paragraph
                MyRun = block_item.runs[i]
                MyRun.font.name = MyFontName
                MyRun.font.size = MyFontSize
                MyRun.font.bold = MyFontBold
                MyRun.font.italic = MyFontItalic
                MyRun.font.underline = MyFontUnderline
                MyRun.font.color.rgb = MyFontColor

        if Text_to_delete2 =="" and re.search(Text_to_delete, block_item.text, flags=0)!= None :# if after run treatment, the text to delete not deleted (runs cut le texte to delete in 2)
            # Try to save the former format of paragraph by saving the format of the last run
            NbRuns = block_item.runs.__len__()
            MyFontName = block_item.runs[NbRuns-1].font.name
            MyFontSize = block_item.runs[NbRuns-1].font.size
            MyFontBold = block_item.runs[NbRuns-1].font.bold
            MyFontItalic = block_item.runs[NbRuns-1].font.italic
            MyFontUnderline = block_item.runs[NbRuns-1].font.underline
            MyFontColor = block_item.runs[NbRuns-1].font.color.rgb
            block_item.text = block_item.text.replace(Text_to_delete, "") # suppress the Text_to_delete but loose the initial look & feel (format) of the paragraph)
            # Try to re establish the former style
            NbRuns = block_item.runs.__len__()
            for i in range(NbRuns):  # Loop through all runs in the paragraph
                MyRun = block_item.runs[i]
                MyRun.font.name = MyFontName
                MyRun.font.size = MyFontSize
                MyRun.font.bold = MyFontBold
                MyRun.font.italic = MyFontItalic
                MyRun.font.underline = MyFontUnderline
                MyRun.font.color.rgb = MyFontColor


    return
#@@@@@@@@@@@@@@@@@@@@@@@@@@@@ END OF "DELETE TEXT IN ONE PARAGRAPH IN FULL TEXT (NO TABLE)" FUNCTION @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@



#@@@@@@@@@@@@@@@@@@@@@@@@@@@@ "INSERT TEXT IN ONE CELL OF A TABLE" FUNCTION @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@
def Insert_Text_Cell (tableCell, TextStart, TextEnd):
    """
    This function inserts text in a cell of a table of Word docx at the beginning and at the end of the text
    This function works only for paragraphs inside tables
    To insert text in paragraphs in full text, another code is required
    This code allows to insert the text into the paragraph using "replace" function which is the only way to do it
    without loosing the initial look & feel of the texte (size, font, color,..)
    because any other way of changing the text of a paragraph in docx Word will unfortunately loose all of that

    Args:
        tableCell : the cell of a table in which we insert the text
        TextStart : the tag to be inserted at the beginning of the paragraph
        TextEnd : the tag to be inserted at the end of the paragraph
            
    Returns:
        The function returns nothing
        but modifies the paragraph in the table cell by adding text
    """
    # scan the paragraphs of the cell and insert the text
    ListOfRuns = []
    for paragCell in tableCell.paragraphs:
        ListOfRuns.extend(paragCell.runs)
    NbRuns = len(ListOfRuns)
    
    if NbRuns == 0: # if the cell has no run
        tableCell.text = tableCell.text.replace(tableCell.text, TextStart + tableCell.text + TextEnd) # insert the TextStart and TextEnd
    else: # if there is at least 1 run
        # insert the start text
        ListOfRuns[0].text = ListOfRuns[0].text.replace("", TextStart,1) 
        # insert the end text
        ListOfRuns[NbRuns-1].text = ListOfRuns[NbRuns-1].text.replace(ListOfRuns[NbRuns-1].text, ListOfRuns[NbRuns-1].text + TextEnd,1)
#@@@@@@@@@@@@@@@@@@@@@@@@@@@@ END OF "INSERT TEXT IN ONE CELL OF A TABLE" FUNCTION @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@



#@@@@@@@@@@@@@@@@@@@@@@@@@@@@ "DELETE TEXT IN ONE CELL OF A TABLE" FUNCTION @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@
def Delete_Text_Cell (tableCell, Text_to_delete):
    """
    This function deletes text in a cell of a table of Word docx
    This function works only for paragraphs inside tables
    To delete text in paragraphs in full text, another code is required
    This code allows to delete the text of the paragraph using "replace" function which is the only way to do it
    without loosing the initial look & feel of the texte (size, font, color,..)
    because any other way of changing the text of a paragraph in docx Word will unfortunately loose all of that

    Args:
        tableCell : the cell of a table in which we delete the text
        Text_to_delete : the text to be deleted in the cell
            
    Returns:
        The function returns nothing
        but modifies the paragraph in the table cell by deleting text
    """        
    
    Text_to_delete2 =""
    # scan the paragraphs of the cell
    ListOfRuns = []
    for paragCell in tableCell.paragraphs:
        ListOfRuns.extend(paragCell.runs)
    NbRuns = len(ListOfRuns)

    if NbRuns == 0: # if the cell has no run
        tableCell.text = tableCell.text.replace(Text_to_delete, "") # replace by ''

    else: # if the paragraph has at least 1 run, we manage at run level
        # we have to create Text_to_delete2 because re.search will not work with simple ? or > or < or /
        if Text_to_delete =="??": 
            Text_to_delete2 = r'\?\?'
        if Text_to_delete =="<>":
            Text_to_delete2 = r'\<\>'
        if Text_to_delete =="</>":
            Text_to_delete2 = r'\<\/\>'

        for i in range(NbRuns):  # Loop through all runs in the paragraph
            MyRun = ListOfRuns [i]
            if Text_to_delete2 !="": # if it is "??" or "<>"" or "</>", we use Text_to_delete2
                if re.search(Text_to_delete2, MyRun.text, flags=0)!= None :
                    MyRun.text = MyRun.text.replace(Text_to_delete, '',1) 
            else: # if it is NOT "??" or "<>"" or "</>", use Text_to_delete
                if re.search(Text_to_delete, MyRun.text, flags=0)!= None :
                    MyRun.text = MyRun.text.replace(Text_to_delete, '',1) 


        if Text_to_delete2 !="" and re.search(Text_to_delete2, tableCell.text, flags=0)!= None :# if after run treatment, the text to delete not deleted (runs cut le texte to delete in 2)
            # Try to save the former format of paragraph by saving the format of the last run
            ListOfRuns = []
            for paragCell in tableCell.paragraphs:
                ListOfRuns.extend(paragCell.runs)
            NbRuns = len(ListOfRuns)
            MyFontName = ListOfRuns [NbRuns-1].font.name
            MyFontSize = ListOfRuns [NbRuns-1].font.size
            MyFontBold = ListOfRuns [NbRuns-1].font.bold
            MyFontItalic = ListOfRuns [NbRuns-1].font.italic
            MyFontUnderline = ListOfRuns [NbRuns-1].font.underline
            MyFontColor = ListOfRuns [NbRuns-1].font.color.rgb
            tableCell.text = tableCell.text.replace(Text_to_delete, "") # suppress the Text_to_delete but loose the initial look & feel (format) of the paragraph)
            # Try to re establish the former style
            ListOfRuns = []
            for paragCell in tableCell.paragraphs:
                ListOfRuns.extend(paragCell.runs)
            NbRuns = len(ListOfRuns)
            for i in range(NbRuns):  # Loop through all runs in the paragraph
                MyRun = ListOfRuns[i]
                MyRun.font.name = MyFontName
                MyRun.font.size = MyFontSize
                MyRun.font.bold = MyFontBold
                MyRun.font.italic = MyFontItalic
                MyRun.font.underline = MyFontUnderline
                MyRun.font.color.rgb = MyFontColor

        if Text_to_delete2 =="" and re.search(Text_to_delete, tableCell.text, flags=0)!= None :# if after run treatment, the text to delete not deleted (runs cut le texte to delete in 2)
            # Try to save the former format of paragraph by saving the format of the last run
            ListOfRuns = []
            for paragCell in tableCell.paragraphs:
                ListOfRuns.extend(paragCell.runs)
            NbRuns = len(ListOfRuns)
            MyFontName = ListOfRuns [NbRuns-1].font.name
            MyFontSize = ListOfRuns [NbRuns-1].font.size
            MyFontBold = ListOfRuns [NbRuns-1].font.bold
            MyFontItalic = ListOfRuns [NbRuns-1].font.italic
            MyFontUnderline = ListOfRuns [NbRuns-1].font.underline
            MyFontColor = ListOfRuns [NbRuns-1].font.color.rgb
            tableCell.text = tableCell.text.replace(Text_to_delete, "") # suppress the Text_to_delete but loose the initial look & feel (format) of the paragraph)
            # Try to re establish the former style
            ListOfRuns = []
            for paragCell in tableCell.paragraphs:
                ListOfRuns.extend(paragCell.runs)
            NbRuns = len(ListOfRuns)
            for i in range(NbRuns):  # Loop through all runs in the paragraph
                MyRun = ListOfRuns[i]
                MyRun.font.name = MyFontName
                MyRun.font.size = MyFontSize
                MyRun.font.bold = MyFontBold
                MyRun.font.italic = MyFontItalic
                MyRun.font.underline = MyFontUnderline
                MyRun.font.color.rgb = MyFontColor
#@@@@@@@@@@@@@@@@@@@@@@@@@@@@ END OF "DELETE TEXT IN ONE CELL OF A TABLE" FUNCTION @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@



#@@@@@@@@@@@@@@@@@@@@@@@@@@@ "READ QUESTIONS AND SIZE ANSWER REQUIREMENTS IN NEW AAP" FUNCTION @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@
def Read_Questions_in_docx ( PathFolderSource, PathForOutputsAndLogs, list_of_SizeWords_OK, list_of_SizeWords_KO, TagQStart = "<>", TagQEnd = "</>" ):
    """
    CONTEXT:
    Uses python-docx 1.1.2 to manipulate Word documents : .docx only but not .doc. You need first to type "pip install python-docx" in your terminal
    Read the questions inside files contained in a folder with .docx extension and which are AAP ("Appel A Projet")
    AAP = document emitted by a donor describing the conditions under which it will grant funds to NGOs

    ACTIONS OF THE CODE
    Finds questions inside the AAP document and finds also information about 
    the size of answer required by the donor, if indicated (not always required)
    (e.g. : number max or min of words, characters, lines,..)
    Puts questions and size requirements into a dictionary with a Unique ID (UID) associated with each question
    Puts the UID into the AAP document at the right place


    ===========        How does it idenfiy the size of answer required ?   ==================
    It uses the lists "list_of_Answer_SizeWords" and "list_of_Exclus_SizeWords"
    In a paragraph, if a word of list_of_Answer_SizeWords is present and no word is present from list_of_Exclus_SizeWords,
    , then this paragraph includes the indication of size. 
    
    The indication of size is generally inside parentesis () in the same paragraph as the question, 
    In this case, it separates the question from the size requirement.
    The indication of size can also be in the following paragraph or a paragraph nearby
    If not found inside the question, it tries to find it around.
    It is common also that there is no indication of size required

    Args:
        PathFolderSource: Path to the folder containing the files to be read
        PathForOutputsAndLogs: Path to the folder containing the log file
        list_of_Answer_SizeWords and list_of_Exclus_SizeWords: used to identify in the texte the requirements 
            for the size of the answer given by the donor (see explainations above)
        TagQStart = "<>" Tag indicating the beginning of a Multi-paragraphs question (question with context below)
        TagQEnd = "</>" Tag indicating the end of a Multi-paragraphs question (question with context below)     

    Returns:
        The function returns a list of dictionaries about the question + size requirement if any in the following format :
[
            {
            ‘uid’: ‘uid’,
            “question”: texte question1 ...
            “size_answer”: size question1
            “enhanced_question”: ...,
            “question_is_open”: close/open
            “question_on_asso”: yes/no
            “response”: .....
            },
            {
            ‘uid’: ‘uid’,
            “question”: texte question2
            “size_answer”: size question2
            “enhanced_question”: ...,
            “question_is_open”: close/open
            “question_on_asso”: yes/no
            “response”: .....
            }, .....
        ]        In the future, we could also add 2 fields in the list of dictionaries for  "general context of AAP", "context for the question"
        Context is not managed for the moment. The field "answer given by IA" will be filled by IA.
        It creates also a new version of the AAP document named "NameOfDocument-with_UID"
        It also logs errors in a file named "logs-IA_for_Asso.txt" in the folder "PathForOutputsAndLogs"
    """
    TheText = '' # Text on which we are working
    DictQuestions = {} #initialise the dictionnary of questions
    ListDict = [] #initialize a list of dictionnaries on questions, in which we put all the info about a single question
    EverythingOK = True # All the prerequisit chekings are OK if True
    #Create a list of path to all the files (no hidden files) contained in the folder “PathFolderSource” 
    FilesWithPath = []
    Multi_Paragraph = False # True if we are in a multi-paragraph,
    Go_DictionUID = False # True if it is OK to send the question to the dictionary and the AAP
    Text_Question = '' # The text of the question that we are going to put in the dictionary

    #======= A - FILES PREREQUISIT CHECK  ========

    for file in glob.glob(PathFolderSource +'*.*'):
        FilesWithPath.append(file)

     #======= A.1 Only docx files in the folder ========
    for file in FilesWithPath:
        TheExtension = file [-4:]
        if TheExtension != "docx":
            EverythingOK = False
            MessageError = str(datetime.now()) + ' Error encountered when reading files : There should only be docx files in the folder'
            logging.error(MessageError)
            print (MessageError)

    for file in FilesWithPath:
        #======= B - OPEN THE DOCX FILE  ========
        NameOfWorkDocument = (file.split('/')[-1])[:-5]
        if EverythingOK and NameOfWorkDocument[len(NameOfWorkDocument)-9:] !="-with UID": #do not put UID where there are already UIDs
            try:
                docWork = Document(file)
            except IOError:
                MessageError = str(datetime.now()) + ' Error encountered when reading Word docx file ' + file
                logging.error(MessageError)
                print(MessageError)

            #============== C - RETRIEVE QUESTIONS AND SIZE REQUIREMENTS ================================================    
            for block_item in iter_block_items(docWork): # scan of working document top down (paragraphs @ tables)

            #============== 1 - RETREIVE QUESTIONS & SIZE IN A "FULL TEXT" PARAGRAPH (NOT IN A TABLE)================================================    

                #=======  1.2 - RETREIVE THE QUESTION FROM THE PARAGRAPH ========
                if isinstance(block_item, Paragraph): # treatment of a "full text" paragraph (not table) 

                    #=======  1.2.1 - TREATMENT IF NOT IN A MULTI-PARAGRAPH ========
                    if Multi_Paragraph == False:
                        if re.search(TagQStart, block_item.text, flags=0) == None : # if there is no TagStart of a multi-paragraph question in it

                            # If "?" in the simple paragraph, it is a question to be added
                            if re.search(r'\?', block_item.text, flags=0)!= None : # if there is a "?" in it
                                Text_Question = block_item.text # put the text into the Text of Question
                                Go_DictionUID = True # Go to send the question to the dictionary and the AAP
                            #else: # There is a tag start but no ?

                    
                        #=======  1.2.2 - TREATMENT IF IT IS THE START OF A MULTI-PARAGRAPH ========
                        else: # if a tag start of multi-paragraph is found = we enter in a Multi-Paragraph
                            Multi_Paragraph = True # it is the start of Multi Paragraph
                            Text_Question = Text_Question + block_item.text # put the text into the Text of Question
                            if re.search(TagQEnd, block_item.text, flags=0) != None : # if there is a TagEnd of a multi-paragraph question in it
                                # there is a Tag start and a Tag End in the same paragraph (should not happen theorically)
                                Multi_Paragraph = False # it is the end of Multi Pragraph
                                Go_DictionUID = True # Go to send the question to the dictionary and the AAP

                        #=======  1.2.3 - TREATMENT IF WE ARE ALREADY IN A MULTI-PARAGRAPH ========
                    else: # Multipragraph = True => we are already in a Multi-Paragraph
                        Text_Question = Text_Question + ' ' + block_item.text # Add the text into the Text of Question
                        if re.search(TagQEnd, block_item.text, flags=0) != None : # if there is a TagEnd of a multi-paragraph question in it
                            Multi_Paragraph = False # it is the end of Multi Pragraph
                            Go_DictionUID = True # Go to send the question to the dictionary and the AAP


                    #=======  1.2.4 - TREATMENT IF IT IS OK TO SEND QUESTION TO DICTIONARY AND AAP ========
                    if Go_DictionUID == True: # If OK to send the question to Dictionary and AAP
                        DictQuestions = {}
                        #ListDict.clear() # Empty the list to put only a part of the paragraph in the question field instead of all the paragraph
                        DictQuestions ["uid"] = uuid.uuid4().hex
                        QuestionUI = DictQuestions ["uid"]


                        #=======  1.2.4.1 - RETREIVE THE SIZE OF ANSWER FROM THE PARAGRAPH (IF ANY) ========

                        #=======  CHECK IF THE TEXT OF THE QUESTION CONTAINS A REQUIREMENT FOR SIZE OF ANSWER AND IF YES, RETRIEVE IT =======
                        if OneOfTheWords_Is_InTheParagraph (Text_Question, list_of_SizeWords_OK, list_of_SizeWords_KO):
                        #======= Manage case of parenthesis in the text = probably a size requirement inside the parenthesis
                            if re.search(r'\(', Text_Question, flags=0)!= None : 
                            # A parenthesis in the text of the paragraph = good probability that it is for size of answer requirement
                                #======= extract the size information if it is inside the parenthesis
                                PosiStart = re.search(r'\(', Text_Question, flags=0).start() # start poition of '('
                                PosiEnd = re.search(r'\)', Text_Question, flags=0).start()+1 # end position of ')'
                                TheText = Text_Question[PosiStart+1:PosiEnd-1] # extract the size information which is betewwen the parenthesis and erase the parenthesis
                                if OneOfTheWords_Is_InTheParagraph (TheText, list_of_SizeWords_OK, list_of_SizeWords_KO):
                                # if there is size info in the paragraph, we must split it into Question + size info
                                    DictQuestions ["question"] = Text_Question[:PosiStart]+' '+Text_Question[PosiEnd:]
                                    DictQuestions ["size_answer"] = TheText
                                else:# Finally, we thought it was a size requirement but we were wrong and it was not
                                    DictQuestions ["question"] = Text_Question
                                    DictQuestions ["size_answer"] = ''
                            else:# No "(" so size requirement is empty)
                                DictQuestions ["question"] = Text_Question
                                DictQuestions ["size_answer"] = ''

                        else: # if no size of answer, put empty size
                            DictQuestions ["question"] = Text_Question
                            DictQuestions ["size_answer"] = ''
                                                        
                        #======= 1.2.4.2 FILL THE EMPTY FIELDS OF THE LIST ==========
                        DictQuestions["enhanced_question"] = ''
                        DictQuestions["question_is_open"] = ''
                        DictQuestions ["question_on_asso"] = ''
                        DictQuestions ["response"] = ''

                        #======= 1.2.4.3 INSERT THE UID AT THE RIGHT PLACE (below the question) =====
                        Insert_Text_Paragraph (block_item, '' , '\n' + QuestionUI)
                        # We use the function Insert_Text_Paragraph to insert '\n' + QuestionUI at the end without loosing the look & feel of the paragraph

                        #======= 1.2.4.4 ADD THE DICTIONARY OF THE QUESTION INTO THE LIST OF DICTIONARIES ======
                        new_dict = DictQuestions.copy() 
                        ListDict.append ( new_dict ) #add the dict to the list of dictionaries
                        DictQuestions.clear() # Empty the dictionary for next question
                        Go_DictionUID = False 
                        Text_Question = '' # Clear the variable Texte_Question



            #============== 2 - RETREIVE QUESTIONS & SIZE IN A CELL OF A TABLE ================================================    

                #=======  2.1 - RETREIVE THE QUESTIONS FROM THE TABLE CELLS  ========
                elif isinstance(block_item, Table): # treatment of a table with cells
                    for row in range(len(block_item.rows)): # Loop on all cells = all rows and all columns
                        for col in range(len(block_item.columns)): #questions are generally in the 1st column but we check "?" everywhere in the table
                            if block_item.cell(row, col).text.strip() != '': # if the cell is not empty

                                #  If "?" in the cell, it is a question to be tagged
                                if re.search(r'\?', block_item.cell(row, col).text, flags=0)!= None : # if there is a "?" in it
                                    DictQuestions.clear() # Empty the dictionary for next question
                                    DictQuestions ["uid"] = uuid.uuid4().hex
                                    DictQuestions ["question"] = block_item.cell(row, col).text
                                    below_is_a_size_for_response = False #Tag indicating that the cell below not empty is a size requirement for the answer


                                    #=======  2.2 - RETREIVE THE SIZE OF ANSWER (IF ANY) FROM THE TABLE CELLS IF A QUESTION HAS BEEN FOUND ========

                                    #=======  CHECK IF THE CELL CONTAINS A REQUIREMENT FOR SIZE OF ANSWER AND IF YES, RETRIEVE IT =======
                                    if OneOfTheWords_Is_InTheParagraph (block_item.cell(row, col).text, list_of_SizeWords_OK, list_of_SizeWords_KO):
                                    #======= Manage case of parenthesis in the text = probably a size requirement inside the parenthesis
                                        if re.search(r'\(', block_item.cell(row, col).text, flags=0)!= None : 
                                        # A parenthesis in the text of the paragraph = good probability that it is for size of answer requirement
                                            #======= extract the size information if it is inside the parenthesis
                                            PosiStart = re.search(r'\(', block_item.cell(row, col).text, flags=0).start() # start poition of '('
                                            PosiEnd = re.search(r'\)', block_item.cell(row, col).text, flags=0).start()+1 # end position of ')'
                                            TheText = block_item.cell(row, col).text[PosiStart+1:PosiEnd-1] # extract the size information which is betewwen the parenthesis and erase the parenthesis
                                            if OneOfTheWords_Is_InTheParagraph (TheText, list_of_SizeWords_OK, list_of_SizeWords_KO):
                                            # if there is size info in the paragraphe, we must split it into Question + size info
                                                DictQuestions.clear() # Empty the dictionary to put only a part of the paragraph in the question field instead of all the paragraph
                                                DictQuestions ["uid"] = uuid.uuid4().hex
                                                DictQuestions ["question"] = block_item.cell(row, col).text[:PosiStart] + ' ' + block_item.cell(row, col).text[PosiEnd:]
                                                DictQuestions ["size_answer"] = TheText

                                            else:# Finally, we thought it was a size requirement but we were wrong and it was not
                                                DictQuestions ["size_answer"] = ''
                                        else:# No "("" so size requirement is empty
                                                DictQuestions ["size_answer"] = ''

                                    elif col < len(block_item.columns)-1 or row < len(block_item.rows)-1 :#if there is a cell at the right or below or above the current cell 
                                        if col < len(block_item.columns)-1:# if cell at the right of the current cell
                                            if OneOfTheWords_Is_InTheParagraph (block_item.cell(row, col+1).text, list_of_SizeWords_OK, list_of_SizeWords_KO):
                                            # test if the size requirement can be in the cell at the right of the current cell 
                                                DictQuestions ["size_answer"] = block_item.cell(row, col+1).text # if key words of size found, we think that it is a size requirement
                                            else:
                                                DictQuestions ["size_answer"] = ''

                                        if row < len(block_item.rows)-1:# if there is a cell below the current cell
                                            if OneOfTheWords_Is_InTheParagraph (block_item.cell(row+1, col).text, list_of_SizeWords_OK, list_of_SizeWords_KO):
                                            # test if the size requirement can be in the cell below of the current cell if there is a cell below
                                                if re.search(r'\(', block_item.cell(row+1, col).text, flags=0).start() == 0: # if '(' is at the 1st position, we think very probable that the cell below is a size requirement
                                                    DictQuestions ["size_answer"] = block_item.cell(row+1, col).text # if key words of size found, we think that it is a size requirement 
                                                    below_is_a_size_for_response = True # we indicate here that the cell below is just a size requirement and that the UID can be put there
                                                else:  
                                                    DictQuestions ["size_answer"] = ''
                                     
                #!!!!! amélioration souhaitable dans la cellule du haut ou du dessous ou de côté, voir si parenthèses et si oui, rechercher si size dans parenthèse et ne prendre que ça, car actuellement on met tout dans size_answer
                                    else: # no size found inside or at the right or below de cell => empty size
                                        DictQuestions ["size_answer"] = ''




                                    #======= 2.3 FILL THE EMPTY FIELDS OF THE LIST ==========
                                    DictQuestions ["enhanced_question"] = ''
                                    DictQuestions ["question_is_open"] = ''
                                    DictQuestions ["question_on_asso"] = ''
                                    DictQuestions ["response"] = ''

                                    #======= 2.4 INSERT THE UID AT THE RIGHT PLACE (below the question) =====
                                    QuestionUI = DictQuestions ["uid"]
                                    #======= Case of Table with only 1 column =====
                                    if len (block_item.columns) == 1 and len (block_item.rows) > row+1: # If there is only one column and there is a cell below
                                        if  block_item.cell(row+1, col).text.strip() == '' or below_is_a_size_for_response == True: # if the cell below it is empty
                                            block_item.cell(row+1, col).text = block_item.cell(row+1, col).text + QuestionUI 
                                            # put the UID in the empty cell below the current one (no insert so the format can be modified)
                                        else: # else, put the UID at the end of the current cell
                                            Insert_Text_Cell (block_item.cell(row, col), '' , '\n' + QuestionUI)
                                    #======= Case of Table with more than 1 column =====
                                    elif len (block_item.columns) > 1 : # If there is more than one column 
                                        if len (block_item.columns) > col+1 :#if there is a column at the right of the current column and
                                            if block_item.cell(row, col+1).text.strip() == '': #and if it is empty
                                                block_item.cell(row, col+1).text = block_item.cell(row, col+1).text + QuestionUI 
                                                # put the UID in the column at the right of the current one (no insert so the format could be modified)
                                            else: # else, put the UID at the end of the current cell
                                                Insert_Text_Cell (block_item.cell(row, col), '' , '\n' + QuestionUI)
                                                # We use the function Insert_Text_Cell to insert '\n' + QuestionUI at the end without loosing the look & feel of the paragraph
                                        else: # else, put the UID at the end of the current cell
                                            Insert_Text_Cell (block_item.cell(row, col), '' , '\n' + QuestionUI)
                                    # We use the function Insert_Text_Cell to insert '\n' + QuestionUI at the end without loosing the look & feel of the paragraph

                                    #======= 2.5 ADD THE UID + QUESTION & INFO INTO THE DICTIONARY ======
                                    new_dict = DictQuestions.copy() 
                                    #new_list = ListDict.copy() 
                                    ListDict.append ( new_dict ) #add the dict to the list of dictionaries
                                    #DictQuestions[QuestionUI] = new_list 
                                    #ListDict.clear() # Empty the list for next question
                                    DictQuestions.clear() # Empty the dictionary for next question

            print(ListDict)
            docWork.save(PathFolderSource + r'/' + NameOfWorkDocument + '-with UID.docx')
        else:
            MessageError = str(datetime.now()) + ' Error encountered when reading Word docx file , please check type .docx and name of the file with no UID)' 
            logging.error(MessageError)
            print(MessageError)

    
    print('End of the read program')



    return ListDict
#@@@@@@@@@@@@@@@@@@@@@@@ END OF "READ QUESTIONS AND SIZE ANSWER REQUIREMENTS IN NEW AAP" FUNCTION @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@



#@@@@@@@@@@@@@@@@@@@@@@@@@@@@ "WRITE ANSWERS INTO NEW AAP AND Q&A DOCX FILE" FUNCTION @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@
def Write_Answers_in_docx (List_UIDQuestionsSizeAnswer, PathFolderSource,  PathForOutputsAndLogs, TagQStart = "<>", TagQEnd = "</>" ):
    """
    CONTEXT:
    Writes answers below or near each question in a docx file AAP ("Appel A Projet")
    Questions and answers associated with Unique IDs (UID) are received in a dictionary 
    An AAP file with de questions associated to UID has already been created by the read function 
        and will be used by the current Write function
    ACTIONS OF THE CODE
    Finds UID inside the AAP document and replace it by the answer associated with the UID from the dictionary

    Args:
        List_UIDQuestionsSizeAnswer: List of dictionaries, each containing the UID + question + Answer
        PathFolderSource: Path to the folder containing the files to be read
        PathForOutputsAndLogs: Path to the folder containing the log file
        TagQStart = "<>" Tag indicating the beginning of a Multi-paragraphs question (question with context below)
        TagQEnd = "</>" Tag indicating the end of a Multi-paragraphs question (question with context below) 

    Returns:
        The function returns nothing but creates 2 files 
        1 file is the AAP with the answers inside associated with the corresponding questions
        1 file is a simple docx file containing questions and answers
        It also logs errors in a file named "logs-IA_for_Asso.txt" in the folder "PathForOutputsAndLogs"
    """

    FilesWithPath = []
    for file in glob.glob(PathFolderSource +'*.*'):
        FilesWithPath.append(file)

    for file in FilesWithPath:
        TheExtension = file [-4:] 
        match TheExtension:
            case 'docx':
                try:
                    f = open(file, 'rb')
                    document = Document(f)
                    NameOfDocument = file.split('/')[-1] # Name of the file without the path will be used in the Key of the dictionnary            
                    
                    if re.search(r'UID', NameOfDocument, flags=0)!= None : # if there is a "UID" in the name of the file
                    # Here, we want to open the file AAP with UID, so we only work with a file including "UID" in its name

                        # for each key of the dictionary, corresponding to the document
                        # find the key in the document and replace it by the answer near the question

                        #============== 1 - SUPPRESS "??", tags "<>" and "</>" IN A "FULL TEXT" PARAGRAPH (NOT IN A TABLE)================================================    
                        for docpara in document.paragraphs:
                            if "??" in docpara.text: # Suppress "??" added to identify questions
                                Delete_Text_Paragraph (docpara, "??")
                                # Do it before inserting the answer to avoid modifying the answer
                            if TagQStart in docpara.text: # Suppress TagQStart added to identify questions
                                Delete_Text_Paragraph (docpara, TagQStart)
                                # Do it before inserting the answer to avoid modifying the answer
                            if TagQEnd in docpara.text: # Suppress TagQEnd added to identify questions
                                Delete_Text_Paragraph (docpara, TagQEnd)
                                # Do it before inserting the answer to avoid modifying the answer

                        #============== 2 - REPLACE UID BY THE ANSWER IN A "FULL TEXT" PARAGRAPH (NOT IN A TABLE)================================================    
                        # Now, we replace the UID keys by the answers in the full text of the document
                        for docpara in document.paragraphs:
                            for value in List_UIDQuestionsSizeAnswer :
                                if value ["uid"] in docpara.text: # key is the UID
                                    Insert_Text_Paragraph (docpara, "" , '\n' + value ["response"] )# Insert the answer
                                    # Suppress key because we have already inserted the answer
                                    Delete_Text_Paragraph (docpara, value ["uid"])

                        #============== 3 - SUPPRESS "??", tags "<>" and "</>" IN A CELL OF A TABLE)================================================    
                        for index, table in enumerate(document.tables):
                            for row in range(len(table.rows)):
                                for col in range(len(table.columns)):
                                    if "??" in table.cell(row, col).text: # Suppress "??" added to identify questions
                                        #table.cell(row, col).text = table.cell(row, col).text.replace("??", "") # suppress the UID
                                        Delete_Text_Cell (table.cell(row, col), "??")
                                    if TagQStart in table.cell(row, col).text: # Suppress TagQStart added to identify questions
                                        #table.cell(row, col).text = table.cell(row, col).text.replace(TagQStart, "") # suppress the UID
                                        Delete_Text_Cell (table.cell(row, col), TagQStart)
                                    if TagQEnd in table.cell(row, col).text: # Suppress TagQEnd added to identify questions
                                        #table.cell(row, col).text = table.cell(row, col).text.replace(TagQEnd, "") # suppress the UID
                                        Delete_Text_Cell (table.cell(row, col), TagQEnd)

                        #============== 4 - REPLACE UID BY THE ANSWER IN A CELL OF A TABLE  ================================================    
                        # then, we replace the keys by the answers in the tables of the document
                        for index, table in enumerate(document.tables):
                            for value in List_UIDQuestionsSizeAnswer :
                                for row in range(len(table.rows)):
                                    for col in range(len(table.columns)):
                                        if value ["uid"] in table.cell(row, col).text:
                                            #table.cell(row, col).text = table.cell(row, col).text.replace(key, value)
                                            Insert_Text_Cell (table.cell(row, col),  "" ,  value ["response"] )# Insert the answer
                                            Delete_Text_Cell (table.cell(row, col), value ["uid"])
                                            #table.cell(row, col).text = table.cell(row, col).text.replace(key, "")# suppress the UID


                        print("==========    DICTIONNAIRE AVEC ANSWERS :   ======")
                        print(List_UIDQuestionsSizeAnswer)

                        #============== 3 - CREATE AAP WITH ANSWERS FILE ================================================    
                        # We create a new version of the AAP document with the answers
                        document.save(PathForOutputsAndLogs+ r'/' + NameOfDocument[:-13] + "_with_answers" + '_' + str(datetime.now())[:-16] + '-' +  str(datetime.now())[-15:-13]+ 'h' + str(datetime.now())[-12:-10]+ 'mn'+ str(datetime.now())[-9:-7]+ 's'+ '.docx' )

                        #============== 4 - CREATE A SIMPLE DOCX FILE WITH QUESTIONS AND ANSWERS ================================================    
                        # We create a new document containing only the questions and answers
                        documentQA = Document()
                        documentQA.add_heading('List of questions and answers of file : '+ '\n' + NameOfDocument[:-14] + '\n' + 'Time : '+ str(datetime.now())[:-16] + '-' +  str(datetime.now())[-15:-13]+ 'h'+ str(datetime.now())[-12:-10]+ 'mn'+ str(datetime.now())[-9:-7]+ 's' + '\n')

                        for value in List_UIDQuestionsSizeAnswer:
                            p = documentQA.add_paragraph()
                            if "??" in value ["question"]: # Suppress from Value ["question"] "??" added to identify questions
                                value ["question"] = value ["question"].replace("??", "") # suppress "??"
                            if TagQStart in value["question"]: # Suppress TagQStart added to identify questions
                                value["question"] = value["question"].replace(TagQStart, "") # suppress TagQStart
                            if TagQEnd in value["question"]: # Suppress TagQEnd added to identify questions
                                value["question"] = value["question"].replace(TagQEnd, "") # suppress TagQEnd
                            Therun = p.add_run(value["question"])
                            Therun.bold = True
                            Therun.font.color.rgb = RGBColor(255, 0, 0)
                            documentQA.add_paragraph('\n' + value ["response"] + '\n')
                        
                        documentQA.save(Path_where_we_put_Outputs+ r'/' + NameOfDocument[:-14]+ '_Q-A' + '_' + str(datetime.now())[:-16] + '-' +  str(datetime.now())[-15:-13]+ 'h'+ str(datetime.now())[-12:-10]+ 'mn'+ str(datetime.now())[-9:-7]+ 's'+ '.docx' )
                
                except IOError:
                        MessageError = str(datetime.now()) + ' Error encountered when opening for writing the Word docx file ' + file
                        logging.error(MessageError)
                        print(MessageError)
                finally:        
                    f.close()

    print('End of the write program')
    return
#@@@@@@@@@@@@@@@@@@@@@@@@@@@@ END OF "WRITE ANSWERS INTO NEW AAP AND Q&A DOCX FILE" FUNCTION @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@



#@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@ MAIN PROGRAM @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@
# Settings for the path files
Path_where_we_put_Outputs = r'C:\Users\Administrateur\Documents\POC\D4G - IA financement\13_ia_financement\Streamlit Vf\LOG'
Folder_where_the_files_are = r'C:\Users\Administrateur\Documents\POC\D4G - IA financement\13_ia_financement\Streamlit Vf\AAP'

# imports
from docx import Document # import de python-docx
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor

import re
from datetime import datetime # for log file
import logging # for log file
import uuid # for unique ID creation
import glob # for file opening & reading

#activate logging of errors in a txt file
logging.basicConfig(filename=Path_where_we_put_Outputs + r'/logs-IA_for_Asso.txt')

# initialize variables
list_of_SizeWords_OK = [
     " MAX", " MIN", " CARACT", " CHARACT", " LIGNE", " LINE", " SIGN", " PAGE",  " PAS EXC", " NOT EXCEED", " MOTS", " WORDS"
         ]

list_of_SizeWords_KO = [
     " SIGNAT", " MAXIMI", " MONTH", " MOIS", " ANS", " ANNé", " YEAR",  " DAY", " JOUR",
     " DURéE", " DURATION", " IMPACT", " AMOUNT", " MONTANT"
         ]
TagQStart = "<>"
TagQEnd = "</>"


# Read the questions in the files and create a dictionary with questions for IA (questions = where there is a question mark "?")
List_UIDQuestionsSize = Read_Questions_in_docx ( Folder_where_the_files_are, Path_where_we_put_Outputs, list_of_SizeWords_OK, list_of_SizeWords_KO, TagQStart , TagQEnd )
# TODO : Send DictQuestionsSizeAnswers to Streamlit

# For the moment, we create a dictionary of answers with the same keys as the dictionary of questions
# by just taking the question as the answer we just put "ANSWER TO: " + the question
List_UIDQuestionsSizeAnswer = List_UIDQuestionsSize.copy()
for value in List_UIDQuestionsSizeAnswer :
    value ["response"] = "ANSWER TO " + value ["question"]
     
# Write the answers into the docx files just below the questions
Write_Answers_in_docx (List_UIDQuestionsSizeAnswer, Folder_where_the_files_are, Path_where_we_put_Outputs, TagQStart , TagQEnd )
#@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@ END OF MAIN PROGRAM @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@
