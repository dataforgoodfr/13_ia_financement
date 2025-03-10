import re
import os
from typing import List, Dict, Optional
import pdfplumber
import docx
from collections import defaultdict
import nltk
from nltk.tokenize import sent_tokenize
from pydantic import BaseModel, Field, field_validator

try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

# Define Pydantic models for structured data
class Question(BaseModel):
    """Model representing an extracted question with its context."""
    text: str = Field(..., description="The text of the question")
    section: str = Field("General", description="The section containing this question")
    context: Optional[str] = Field(None, description="Surrounding context of the question")
    line_number: Optional[int] = Field(None, description="Line number in the original document")
    
    @field_validator('text')
    def validate_question_text(cls, v):
        """Validate that the question text is not empty."""
        v = v.strip()
        if not v:
            raise ValueError("Question text cannot be empty")
        return v
    
    class Config:
        frozen = True # Prevent object modification

class Section(BaseModel):
    """Model representing a document section."""
    name: str
    questions: List[Question] = Field(default_factory=list)

class DocumentMetadata(BaseModel):
    """Model containing metadata about the processed document."""
    file_path: str
    document_type: str
    total_questions: int = 0
    total_sections: int = 0

class ExtractedDocument(BaseModel):
    """Model representing the complete extracted document."""
    metadata: DocumentMetadata
    sections: Dict[str, Section] = Field(default_factory=dict)
    questions: List[Question] = Field(default_factory=list)
    raw_text: str
    
    def add_question(self, question: Question):
        """Add a question to the document and update counts."""
        self.questions.append(question)
        
        # Add to appropriate section
        if question.section not in self.sections:
            self.sections[question.section] = Section(name=question.section)
        
        # Create a section-specific version of the question
        section_questions = list(self.sections[question.section].questions)
        section_questions.append(question)
        
        # Update the section with the new list of questions
        self.sections[question.section] = Section(
            name=question.section,
            questions=section_questions
        )
        
        # Update metadata
        self.metadata.total_questions = len(self.questions)
        self.metadata.total_sections = len(self.sections)
    
    class Config:
        validate_assignment = True  # Validate attributes when they're assigned
        
class AAPExtractor:
    """
    Class for extracting questions from Word or PDF documents using Pydantic models.
    """
    
    def __init__(self, question_patterns: Optional[List[str]] = None):
        """
        Initialize the AAPExtractor with optional custom patterns.
        
        Args:
            question_patterns: List of regex patterns to identify questions.
                               If None, default patterns will be used.
        """
        # Default patterns to identify questions in various formats
        self.question_patterns = question_patterns or [
            r"(?:^|\n)([0-9]+[\.\)]\s*[A-Z].*?\?)",  # Questions numérotées finissant par un "?""
            r"(?:^|\n)([A-Z][^.?!]+\?)",  # Phrases commençant par une majuscule et finissant par un "?""
            r"(?:^|\n)([Qq]uestion(?:\s+[0-9]+)?[:.]\s*)(.*?)(?=\n|$)",  # Label "Question:" explicite
            r"(?<=\n)([A-Za-z0-9][^.?!]+\?)",  # Toute ligne finissant par un "?""
            r"(?<=\n)([A-Za-z0-9].*?)(?=\n\s*_+|\n\s*\[.*?\]|\n\s*\(.*?\))",  # Texte suivi de underscores ou brackets
            r"(?<=\n)(.*?)(?=:\s*_+|:\s*\[.*?\]|:\s*\(.*?\))",  # Labels suivi de virgules et underscores ou brackets
        ]
        
        # Patterns that indicate fillable areas
        self.fillable_indicators = [
            r"_+",  # Underscores
            r"\[.*?\]",  # Square brackets
            r"\(.*?\)",  # Parentheses with content
            r"□|\u2610",  # Checkbox symbols
        ]
        
        # Words that often indicate a question or prompt
        self.question_indicators_en = [
            "describe", "explain", "list", "define", "compare", "analyze",
            "evaluate", "discuss", "identify", "outline", "summarize",
            "provide", "state", "justify", "elaborate", "specify",
            "expected", "plan", "schedule", "contact", "please",
            "applicant", "candidate", "introducer", "background",
            "activities", "organization", "situation", "address",
            "summary", "impact", "budget", "duration", "financing"
        ]
        
        self.question_indicators_fr = [
            "calendrier", "contact", "actions",  "descriptions", "détaillée", "thématiques",
            "candidat", "introduire", "antécédents", "activités", "organisation", "situation",
            "résumé", "impact", "budget", "durée", "financements", "soutien", "partenaires",
            "Nom du projet", "Objet du financement", "Représentant légal", "Adresse du siège social",
            "Nombre d'adhérents", "Nombre de bénévoles", "Nombre de salariés", "Nombre de personnes bénéficiaires",
            "structure", "écosystème", "Implantation géographique", "Public cible", "Statut juridique",
            "court descriptif", "informations générales"
        ]
        
        # Initialize storage for extracted information
        self.extracted_text = ""
        self.document_path = None
        self.document_type = None
        self.sections = defaultdict(list)
        self.document = None
    
    def extract_from_file(self, file_path: str) -> ExtractedDocument:
        """
        Extract text from file depending on the file type.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ExtractedDocument containing all extracted information
        """
        self.document_path = file_path
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            self.document_type = 'pdf'
            self.extracted_text = self._extract_from_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            self.document_type = 'docx'
            self.extracted_text = self._extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
        
        # Create the Pydantic document model
        metadata = DocumentMetadata(
            file_path=file_path,
            document_type=self.document_type,
            total_questions=0,
            total_sections=0
        )
        
        self.document = ExtractedDocument(
            metadata=metadata,
            raw_text=self.extracted_text
        )
        
        # Process the document to extract questions and sections
        self._process_document()
        
        return self.document
    
    def _extract_from_pdf(self, pdf_path: str) -> str:
        """
        Extract text from pdf.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Extracted text content as a string
        """
        try:
            with pdfplumber.open(pdf_path) as pdf:
                pages_text = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages_text.append(text)
                
                return "\n".join(pages_text)
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return ""
    
    def _extract_from_docx(self, docx_path: str) -> str:
        """
        Extract text from word.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Extracted text content as a string
        """
        try:
            doc = docx.Document(docx_path)
            paragraphs = [p.text for p in doc.paragraphs]
            return "\n".join(paragraphs)
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return ""
    
    def _process_document(self):
        """Process the document to extract questions and identify sections."""
        # First identify sections
        self.identify_sections()
        
        # Extract questions
        questions = self.extract_questions()
        
        # Add questions to the document model
        line_numbers = self._get_question_line_numbers(questions)
        
        for i, question_text in enumerate(questions):
            line_number = line_numbers.get(question_text)
            
            # Find the section
            section = "General"
            for sec_name, sec_questions in self.sections.items():
                if question_text in sec_questions:
                    section = sec_name
                    break
            
            # Create the question model
            question = Question(
                text=question_text,
                section=section,
                line_number=line_number
            )
            
            # Add to document
            self.document.add_question(question)
    
    def _get_question_line_numbers(self, questions: List[str]) -> Dict[str, int]:
        """Get line numbers for each question in the text."""
        result = {}
        lines = self.extracted_text.split('\n')
        
        for question in questions:
            for i, line in enumerate(lines, 1):
                if question in line:
                    result[question] = i
                    break
        
        return result
    
    def preprocess_text(self, text: Optional[str] = None) -> str:
        """
        Clean and preprocess the extracted text.
        
        Args:
            text: Text to preprocess. If None, uses self.extracted_text
            
        Returns:
            Preprocessed text
        """
        if text is None:
            text = self.extracted_text
            
        # Basic cleaning
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n+', '\n', text)
        
        # Standardize question indicators
        for indicator in ["Question", "Q."]:
            text = re.sub(rf'({indicator}\s*[0-9]+)\s*[-:.]', r'\1: ', text, flags=re.IGNORECASE)
        
        # Restore meaningful line breaks
        text = re.sub(r'([.?!])\s+', r'\1\n', text)
        
        return text
    
    def extract_questions(self, text: Optional[str] = None) -> List[str]:
        """
        Extract questions from text using multiple pattern recognition strategies.
        
        Args:
            text: Text to extract questions from. If None, uses self.extracted_text
            
        Returns:
            List of extracted questions
        """
        if text is None:
            text = self.extracted_text
        
        # Clean text before processing
        processed_text = self.preprocess_text(text)
        
        # Storage for found questions using multiple methods
        all_questions = set()
        
        # Method 1: Pattern-based extraction
        for pattern in self.question_patterns:
            matches = re.finditer(pattern, processed_text, re.MULTILINE | re.DOTALL)
            for match in matches:
                if len(match.groups()) > 1:
                    # If the pattern captures both label and content
                    question = match.group(1).strip() + " " + match.group(2).strip()
                else:
                    question = match.group(1).strip()
                
                # Clean up the question
                question = self._clean_question(question)
                if question:
                    all_questions.add(question)
        
        # Method 2: Context-based extraction - find text before fillable areas
        lines = processed_text.split('\n')
        for i, line in enumerate(lines):
            # Skip empty lines
            if not line.strip():
                continue
                
            # Check if the next line contains a fillable indicator
            if i < len(lines) - 1:
                next_line = lines[i + 1]
                is_fillable = any(re.search(pattern, next_line) for pattern in self.fillable_indicators)
                
                if is_fillable:
                    # Current line might be a question/prompt
                    question = self._clean_question(line)
                    if question:
                        all_questions.add(question)
            
            # Check if the current line contains a fillable area
            for pattern in self.fillable_indicators:
                if re.search(pattern, line):
                    # Extract the text before the fillable area
                    before_fillable = re.split(pattern, line)[0].strip()
                    if before_fillable:
                        question = self._clean_question(before_fillable)
                        if question:
                            all_questions.add(question)
        
        # Method 3: Linguistic-based extraction - sentences with question words or indicators
        sentences = sent_tokenize(processed_text)
        for sentence in sentences:
            sentence = sentence.strip()
            # Skip too short sentences
            if len(sentence) < 5:
                continue
                
            # Check if it ends with a question mark
            if sentence.endswith('?'):
                question = self._clean_question(sentence)
                if question:
                    all_questions.add(question)
                continue
                
            # Check for question indicators (like "describe", "explain", etc.)
            lower_sent = sentence.lower()
            if any(lower_sent.startswith(indicator) for indicator in self.question_indicators_en + self.question_indicators_fr):
                question = self._clean_question(sentence)
                if question and not any(q.lower() == question.lower() for q in all_questions):
                    all_questions.add(question)
        
        return list(all_questions)
    
    def _clean_question(self, text: str) -> str:
        """
        Clean a potential question text.
        
        Args:
            text: Text to clean
            
        Returns:
            Cleaned question text
        """
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Remove fillable indicators from the question text
        for pattern in self.fillable_indicators:
            text = re.sub(pattern, '', text)
        
        # Remove common non-question text
        text = re.sub(r'^(Please\s+|Note:\s+)', '', text, flags=re.IGNORECASE)
        
        # Ensure there's actual content
        if len(text) < 5 or not any(c.isalpha() for c in text):
            return ""
        
        return text
    
    def identify_sections(self, text: Optional[str] = None) -> Dict[str, List[str]]:
        """
        Identify document sections and questions within each section.
        
        Args:
            text: Text to process. If None, uses self.extracted_text
            
        Returns:
            Dictionary of sections with their questions
        """
        if text is None:
            text = self.extracted_text
            
        # Try to identify sections based on headers
        section_patterns = [
            r'(?:^|\n)(?:Section|Part)\s+([0-9A-Z]+)[:.]\s*([^\n]+)',
            r'(?:^|\n)([A-Z][A-Z\s]+)(?:\r?\n|\s*:\s*)'
        ]
        
        current_section = "General"
        sections = defaultdict(list)
        
        # First pass: identify sections
        processed_text = self.preprocess_text(text)
        lines = processed_text.split('\n')
        
        for i, line in enumerate(lines):
            # Check if this line is a section header
            is_section_header = False
            for pattern in section_patterns:
                match = re.search(pattern, line)
                if match:
                    if len(match.groups()) > 1:
                        current_section = f"{match.group(1)}: {match.group(2)}"
                    else:
                        current_section = match.group(1)
                    is_section_header = True
                    break
            
            if not is_section_header:
                # Look for questions in this section
                question = self._clean_question(line)
                if question and question.endswith('?'):
                    sections[current_section].append(question)
                    
                # Check if line followed by fillable area in next line
                if i < len(lines) - 1:
                    next_line = lines[i + 1]
                    is_fillable = any(re.search(pattern, next_line) for pattern in self.fillable_indicators)
                    
                    if is_fillable and question:
                        sections[current_section].append(question)
        
        self.sections = dict(sections)
        return self.sections
    
    def get_questions_with_context(self, window_size: int = 1) -> List[Question]:
        """
        Get questions with surrounding context (preceding and following lines).
        
        Args:
            window_size: Number of lines before and after to include as context
            
        Returns:
            List of Question models with context
        """
        if not self.document:
            raise ValueError("No document has been processed yet")
            
        lines = self.extracted_text.split('\n')
        questions_with_context = []
        
        for question in self.document.questions:
            if question.line_number:
                # Get context before and after
                start = max(0, question.line_number - window_size - 1)  # -1 because line_number is 1-indexed
                end = min(len(lines), question.line_number + window_size)
                
                context = '\n'.join(lines[start:end])
                
                # Create a new question with context
                updated_question = Question(
                    text=question.text,
                    section=question.section,
                    context=context,
                    line_number=question.line_number
                )
                
                questions_with_context.append(updated_question)
        
        return questions_with_context
    
    def export_to_json(self, output_file: str) -> str:
        """
        Export the extracted document to a JSON file.
        
        Args:
            output_file: Path to save the JSON
            
        Returns:
            Path to the saved file
        """
        if not self.document:
            raise ValueError("No document has been processed yet")
            
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(self.document.json(indent=2))
            
        return output_file
    
    def export_questions(self, output_file: Optional[str] = None, include_context: bool = False):
        """
        Export extracted questions to a file.
        
        Args:
            output_file: Path to save the questions. If None, returns as string
            include_context: Whether to include context with questions
            
        Returns:
            If output_file is None, returns the formatted questions as a string
        """
        if not self.document:
            raise ValueError("No document has been processed yet")
            
        output = []
        
        if include_context:
            questions_with_context = self.get_questions_with_context()
            for i, question in enumerate(questions_with_context, 1):
                output.append(f"Question {i}: {question.text}")
                output.append(f"Section: {question.section}")
                if question.context:
                    output.append(f"Context: {question.context}")
                output.append("-" * 40)
        else:
            # Group by section
            output.append("Questions by Section:\n")
            for section_name, section in self.document.sections.items():
                output.append(f"Section: {section_name}")
                
                for i, question in enumerate(section.questions, 1):
                    output.append(f"  {i}. {question.text}")
                
                output.append("")  # Empty line between sections
                
            # List all questions in order
            output.append("All Questions (in order of appearance):")
            for i, question in enumerate(self.document.questions, 1):
                output.append(f"{i}. {question.text}")
        
        formatted_output = "\n".join(output)
        
        if output_file:
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(formatted_output)
            return f"Questions saved to {output_file}"
        else:
            return formatted_output
    
    def get_statistics(self) -> Dict:
        """
        Get statistics about the extracted questions.
        
        Returns:
            Dictionary with statistics
        """
        if not self.document:
            raise ValueError("No document has been processed yet")
            
        stats = {
            "document_type": self.document.metadata.document_type,
            "document_path": self.document.metadata.file_path,
            "extraction_date": self.document.metadata.extraction_date,
            "total_questions": self.document.metadata.total_questions,
            "avg_question_length": sum(len(q.text) for q in self.document.questions) / max(1, len(self.document.questions)),
            "sections": len(self.document.sections),
            "questions_per_section": {
                section.name: len(section.questions) 
                for section_name, section in self.document.sections.items()
            }
        }
        
        return stats