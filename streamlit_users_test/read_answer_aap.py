# Préambule 
"""
Format d echange de dictionnaire : 
{
   'uid': 'uid',
   "question": ...
   "enhanced_question": ...,
   "question_is_open": close/open
   "question_on_asso": yes/no
 }
"""
# Import des bibliothèques
import docx
from docx import Document
from docx.document import Document as DocxDocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor
import re
from datetime import datetime 
import logging
import uuid 
import glob 
import os
from pathlib import Path

# Get the directory of the current script (e.g., app.py)
SCRIPT_DIR = Path(__file__).parent.resolve()


# Fonction pour scanner les paragraphes et les tables 
def iter_block_items(parent):
    if isinstance(parent, DocxDocumentType):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


# Fonction pour détecter la présence d'un mot dans une liste pré-définie
def OneOfTheWords_Is_InTheParagraph (TheText, list_of_Words_OK, list_of_Words_KO):
    """
        Args:
        docpara : the paragraph in which we verify
        list_of_Words_OK : List of words that we want to check if they are present in the paragraph
        list_of_Words_KO : List of words that indicate wrong interpretation of the Words OK
    """
    FlagWord_OK = False 
    for Theword in list_of_Words_OK: 
        if re.search(Theword.lower(), TheText.lower(), flags=0)!= None:
            FlagWord_OK = True

    for Theword in list_of_Words_KO: 
        if re.search(Theword.lower(), TheText.lower(), flags=0)!= None:
            FlagWord_OK = False 
    return FlagWord_OK 


# Fonction pour insérer du texte dans un paragraphe
def Insert_Text_Paragraph(block_item, TextStart, TextEnd):
    """
    Args:
        block_item : the paragraph in which we insert the text
        TextStart : the text to be inserted at the beginning of the paragraph
        TextEnd : the text to be inserted at the end of the paragraph
    """
    if block_item.runs == []:
        block_item.text = TextStart + block_item.text + TextEnd 
    else: 
        block_item.runs[0].text = block_item.runs[0].text.replace("", TextStart,1) 
        NbRuns = block_item.runs.__len__()
        block_item.runs[NbRuns-1].text = block_item.runs[NbRuns-1].text.replace(block_item.runs[NbRuns-1].text, block_item.runs[NbRuns-1].text + TextEnd,1)

    return



# Fonction pour supprimer du texte dans un paragraphe
def Delete_Text_Paragraph (block_item, Text_to_delete):
    """
    Args:
        block_item : the paragraph in which we insert the tags
        Text_to_delete : the text to be deleted in the paragraph
    """
    Text_to_delete2 =""
    if block_item.runs == []: 
        block_item.text = block_item.text.replace(Text_to_delete, "") 

    else: 
        if Text_to_delete =="??": 
            Text_to_delete2 = r'\?\?'
        if Text_to_delete =="<>":
            Text_to_delete2 = r'\<\>'
        if Text_to_delete =="</>":
            Text_to_delete2 = r'\<\/\>'
        
        NbRuns = block_item.runs.__len__()
        for i in range(NbRuns):  
            MyRun = block_item.runs[i]
            if Text_to_delete2 !="": 
                if re.search(Text_to_delete2, MyRun.text, flags=0)!= None :
                    MyRun.text = MyRun.text.replace(Text_to_delete, '',1) 
            else: 
                if re.search(Text_to_delete, MyRun.text, flags=0)!= None :
                    MyRun.text = MyRun.text.replace(Text_to_delete, '',1) 
        if Text_to_delete2 !="" and re.search(Text_to_delete2, block_item.text, flags=0)!= None :

            NbRuns = block_item.runs.__len__()
            MyFontName = block_item.runs[NbRuns-1].font.name
            MyFontSize = block_item.runs[NbRuns-1].font.size
            MyFontBold = block_item.runs[NbRuns-1].font.bold
            MyFontItalic = block_item.runs[NbRuns-1].font.italic
            MyFontUnderline = block_item.runs[NbRuns-1].font.underline
            MyFontColor = block_item.runs[NbRuns-1].font.color.rgb
            block_item.text = block_item.text.replace(Text_to_delete, "") 
            NbRuns = block_item.runs.__len__()
            for i in range(NbRuns):  
                MyRun = block_item.runs[i]
                MyRun.font.name = MyFontName
                MyRun.font.size = MyFontSize
                MyRun.font.bold = MyFontBold
                MyRun.font.italic = MyFontItalic
                MyRun.font.underline = MyFontUnderline
                MyRun.font.color.rgb = MyFontColor

        if Text_to_delete2 =="" and re.search(Text_to_delete, block_item.text, flags=0)!= None :
            NbRuns = block_item.runs.__len__()
            MyFontName = block_item.runs[NbRuns-1].font.name
            MyFontSize = block_item.runs[NbRuns-1].font.size
            MyFontBold = block_item.runs[NbRuns-1].font.bold
            MyFontItalic = block_item.runs[NbRuns-1].font.italic
            MyFontUnderline = block_item.runs[NbRuns-1].font.underline
            MyFontColor = block_item.runs[NbRuns-1].font.color.rgb
            block_item.text = block_item.text.replace(Text_to_delete, "")
            NbRuns = block_item.runs.__len__()
            for i in range(NbRuns): 
                MyRun = block_item.runs[i]
                MyRun.font.name = MyFontName
                MyRun.font.size = MyFontSize
                MyRun.font.bold = MyFontBold
                MyRun.font.italic = MyFontItalic
                MyRun.font.underline = MyFontUnderline
                MyRun.font.color.rgb = MyFontColor


    return

# Fonction pour insérer du texte dans un tableau 
def Insert_Text_Cell (tableCell, TextStart, TextEnd):
    """
    Args:
        tableCell : the cell of a table in which we insert the text
        TextStart : the tag to be inserted at the beginning of the paragraph
        TextEnd : the tag to be inserted at the end of the paragraph
    """
    ListOfRuns = []
    for paragCell in tableCell.paragraphs:
        ListOfRuns.extend(paragCell.runs)
    NbRuns = len(ListOfRuns)
    
    if NbRuns == 0:
        tableCell.text = tableCell.text.replace(tableCell.text, TextStart + tableCell.text + TextEnd)
    else: 
        ListOfRuns[0].text = ListOfRuns[0].text.replace("", TextStart,1) 
        ListOfRuns[NbRuns-1].text = ListOfRuns[NbRuns-1].text.replace(ListOfRuns[NbRuns-1].text, ListOfRuns[NbRuns-1].text + TextEnd,1)


# Fonction pour supprimer du texte dans un tableau 
def Delete_Text_Cell (tableCell, Text_to_delete):
    """
    Args:
        tableCell : the cell of a table in which we delete the text
        Text_to_delete : the text to be deleted in the cell
    """        
    
    Text_to_delete2 =""
    ListOfRuns = []
    for paragCell in tableCell.paragraphs:
        ListOfRuns.extend(paragCell.runs)
    NbRuns = len(ListOfRuns)

    if NbRuns == 0:
        tableCell.text = tableCell.text.replace(Text_to_delete, "") 

    else: 
        if Text_to_delete =="??": 
            Text_to_delete2 = r'\?\?'
        if Text_to_delete =="<>":
            Text_to_delete2 = r'\<\>'
        if Text_to_delete =="</>":
            Text_to_delete2 = r'\<\/\>'

        for i in range(NbRuns):  
            MyRun = ListOfRuns [i]
            if Text_to_delete2 !="": 
                if re.search(Text_to_delete2, MyRun.text, flags=0)!= None :
                    MyRun.text = MyRun.text.replace(Text_to_delete, '',1) 
            else: 
                if re.search(Text_to_delete, MyRun.text, flags=0)!= None :
                    MyRun.text = MyRun.text.replace(Text_to_delete, '',1) 


        if Text_to_delete2 !="" and re.search(Text_to_delete2, tableCell.text, flags=0)!= None :
            ListOfRuns = []
            for paragCell in tableCell.paragraphs:
                ListOfRuns.extend(paragCell.runs)
            NbRuns = len(ListOfRuns)
            MyFontName = ListOfRuns [NbRuns-1].font.name
            MyFontSize = ListOfRuns [NbRuns-1].font.size
            MyFontBold = ListOfRuns [NbRuns-1].font.bold
            MyFontItalic = ListOfRuns [NbRuns-1].font.italic
            MyFontUnderline = ListOfRuns [NbRuns-1].font.underline
            MyFontColor = ListOfRuns [NbRuns-1].font.color.rgb
            tableCell.text = tableCell.text.replace(Text_to_delete, "") 
            ListOfRuns = []
            for paragCell in tableCell.paragraphs:
                ListOfRuns.extend(paragCell.runs)
            NbRuns = len(ListOfRuns)
            for i in range(NbRuns):  
                MyRun = ListOfRuns[i]
                MyRun.font.name = MyFontName
                MyRun.font.size = MyFontSize
                MyRun.font.bold = MyFontBold
                MyRun.font.italic = MyFontItalic
                MyRun.font.underline = MyFontUnderline
                MyRun.font.color.rgb = MyFontColor

        if Text_to_delete2 =="" and re.search(Text_to_delete, tableCell.text, flags=0)!= None :
            ListOfRuns = []
            for paragCell in tableCell.paragraphs:
                ListOfRuns.extend(paragCell.runs)
            NbRuns = len(ListOfRuns)
            MyFontName = ListOfRuns [NbRuns-1].font.name
            MyFontSize = ListOfRuns [NbRuns-1].font.size
            MyFontBold = ListOfRuns [NbRuns-1].font.bold
            MyFontItalic = ListOfRuns [NbRuns-1].font.italic
            MyFontUnderline = ListOfRuns [NbRuns-1].font.underline
            MyFontColor = ListOfRuns [NbRuns-1].font.color.rgb
            tableCell.text = tableCell.text.replace(Text_to_delete, "") 
            ListOfRuns = []
            for paragCell in tableCell.paragraphs:
                ListOfRuns.extend(paragCell.runs)
            NbRuns = len(ListOfRuns)
            for i in range(NbRuns):  
                MyRun = ListOfRuns[i]
                MyRun.font.name = MyFontName
                MyRun.font.size = MyFontSize
                MyRun.font.bold = MyFontBold
                MyRun.font.italic = MyFontItalic
                MyRun.font.underline = MyFontUnderline
                MyRun.font.color.rgb = MyFontColor

# Fonction pour lire les questions et la taille des réponses souhaitée
def Read_Questions_in_docx ( 
        PathFolderSource, PathForOutputsAndLogs, list_of_SizeWords_OK, 
        list_of_SizeWords_KO, TagQStart = "<>", TagQEnd = "</>" 
    ):
    """
    Args:
        PathFolderSource: Path to the folder containing the files to be read
        PathForOutputsAndLogs: Path to the folder containing the log file
        list_of_Answer_SizeWords and list_of_Exclus_SizeWords: used to identify in the texte the requirements 
            for the size of the answer given by the donor (see explainations above)
        TagQStart = "<>" Tag indicating the beginning of a Multi-paragraphs question (question with context below)
        TagQEnd = "</>" Tag indicating the end of a Multi-paragraphs question (question with context below)     
    """
    TheText = '' 
    DictQuestions = {} 
    ListDict = [] 
    EverythingOK = True 
    FilesWithPath = []
    Multi_Paragraph = False 
    Go_DictionUID = False 
    Text_Question = ''

    # =========== modifs pour déploiement streamlit
    # Conservation de votre variable originale
    PathFolderSource = SCRIPT_DIR / 'output_aap'

    PathForOutputsAndLogs = SCRIPT_DIR / PathForOutputsAndLogs  # Converti en objet Path

    # Récupération des fichiers avec chemins absolus
    FilesWithPath = [file.resolve() for file in PathFolderSource.glob('*.*')]

    # Affichage des résultats (pour vérification)
    print(FilesWithPath)

    #===============================================

    for file in FilesWithPath:
        file="/".join(file.parts).replace("//", "/")
        TheExtension = file [-4:]
        if TheExtension != "docx":
            EverythingOK = False
            MessageError = str(datetime.now()) + ' Error encountered when reading files : There should only be docx files in the folder'
            logging.error(MessageError)
            print (MessageError)

    for file in FilesWithPath:

        NameOfWorkDocument = os.path.splitext(os.path.basename(file))[0]
        if EverythingOK and NameOfWorkDocument[len(NameOfWorkDocument)-9:] !="-with UID": 
            try:
                docWork = docx.Document(file)
            except IOError:
                MessageError = str(datetime.now()) + ' Error encountered when reading Word docx file ' + file
                logging.error(MessageError)
                print(MessageError)

            for block_item in iter_block_items(docWork): 
                if isinstance(block_item, Paragraph):
                    if Multi_Paragraph == False:
                        if re.search(TagQStart, block_item.text, flags=0) == None : 
                            if '??' in block_item.text or re.search(re.escape(TagQStart), block_item.text):
                                Text_Question = block_item.text 
                                Go_DictionUID = True 

                        else: 
                            Multi_Paragraph = True 
                            Text_Question = Text_Question + block_item.text 
                            if re.search(TagQEnd, block_item.text, flags=0) != None : 
                                Multi_Paragraph = False 
                                Go_DictionUID = True 

                    else: 
                        Text_Question = Text_Question + ' ' + block_item.text 
                        if re.search(TagQEnd, block_item.text, flags=0) != None : 
                            Multi_Paragraph = False 
                            Go_DictionUID = True 

                    if Go_DictionUID == True: 
                        DictQuestions = {}
                        DictQuestions ["uid"] = uuid.uuid4().hex
                        QuestionUI = DictQuestions ["uid"]

                        if OneOfTheWords_Is_InTheParagraph (Text_Question, list_of_SizeWords_OK, list_of_SizeWords_KO):
                            if re.search(r'\(', Text_Question, flags=0)!= None : 
                                PosiStart = re.search(r'\(', Text_Question, flags=0).start() 
                                PosiEnd = re.search(r'\)', Text_Question, flags=0).start()+1 
                                TheText = Text_Question[PosiStart+1:PosiEnd-1] 
                                if OneOfTheWords_Is_InTheParagraph (TheText, list_of_SizeWords_OK, list_of_SizeWords_KO):
                                    DictQuestions ["question"] = Text_Question[:PosiStart]+' '+Text_Question[PosiEnd:]
                                    DictQuestions ["size_answer"] = TheText
                                else:
                                    DictQuestions ["question"] = Text_Question
                                    DictQuestions ["size_answer"] = ''
                            else:
                                DictQuestions ["question"] = Text_Question
                                DictQuestions ["size_answer"] = ''

                        else: 
                            DictQuestions ["question"] = Text_Question
                            DictQuestions ["size_answer"] = ''
                                                        
                        DictQuestions["enhanced_question"] = ''
                        DictQuestions["question_is_open"] = ''
                        DictQuestions ["question_on_asso"] = ''
                        DictQuestions ["response"] = ''

                        Insert_Text_Paragraph (block_item, '' , '\n' + QuestionUI)
                        
                        new_dict = DictQuestions.copy() 
                        ListDict.append ( new_dict ) 
                        DictQuestions.clear() 
                        Go_DictionUID = False 
                        Text_Question = '' 


                elif isinstance(block_item, Table): 
                    for row in range(len(block_item.rows)): 
                        for col in range(len(block_item.columns)): 
                            if block_item.cell(row, col).text.strip() != '': 

                                if re.search(r'\?', block_item.cell(row, col).text, flags=0)!= None : 
                                    DictQuestions.clear() 
                                    DictQuestions ["uid"] = uuid.uuid4().hex
                                    DictQuestions ["question"] = block_item.cell(row, col).text
                                    below_is_a_size_for_response = False 


                                    if OneOfTheWords_Is_InTheParagraph (block_item.cell(row, col).text, list_of_SizeWords_OK, list_of_SizeWords_KO):
                                        if re.search(r'\(', block_item.cell(row, col).text, flags=0)!= None : 
                                            PosiStart = re.search(r'\(', block_item.cell(row, col).text, flags=0).start() 
                                            PosiEnd = re.search(r'\)', block_item.cell(row, col).text, flags=0).start()+1 
                                            TheText = block_item.cell(row, col).text[PosiStart+1:PosiEnd-1] 
                                            if OneOfTheWords_Is_InTheParagraph (TheText, list_of_SizeWords_OK, list_of_SizeWords_KO):
                                                DictQuestions.clear() 
                                                DictQuestions ["uid"] = uuid.uuid4().hex
                                                DictQuestions ["question"] = block_item.cell(row, col).text[:PosiStart] + ' ' + block_item.cell(row, col).text[PosiEnd:]
                                                DictQuestions ["size_answer"] = TheText

                                            else:
                                                DictQuestions ["size_answer"] = ''
                                        else:
                                                DictQuestions ["size_answer"] = ''

                                    elif col < len(block_item.columns)-1 or row < len(block_item.rows)-1 : 
                                        if col < len(block_item.columns)-1:
                                            if OneOfTheWords_Is_InTheParagraph (block_item.cell(row, col+1).text, list_of_SizeWords_OK, list_of_SizeWords_KO):
                                                DictQuestions ["size_answer"] = block_item.cell(row, col+1).text 
                                            else:
                                                DictQuestions ["size_answer"] = ''

                                        if row < len(block_item.rows)-1:
                                            if OneOfTheWords_Is_InTheParagraph (block_item.cell(row+1, col).text, list_of_SizeWords_OK, list_of_SizeWords_KO):
                                                if re.search(r'\(', block_item.cell(row+1, col).text, flags=0).start() == 0:
                                                    DictQuestions ["size_answer"] = block_item.cell(row+1, col).text 
                                                    below_is_a_size_for_response = True 
                                                else:  
                                                    DictQuestions ["size_answer"] = ''
                                     
                                    else: 
                                        DictQuestions ["size_answer"] = ''

                                    DictQuestions ["enhanced_question"] = ''
                                    DictQuestions ["question_is_open"] = ''
                                    DictQuestions ["question_on_asso"] = ''
                                    DictQuestions ["response"] = ''

                                    QuestionUI = DictQuestions ["uid"]
                                    if len (block_item.columns) == 1 and len (block_item.rows) > row+1: 
                                        if  block_item.cell(row+1, col).text.strip() == '' or below_is_a_size_for_response == True: 
                                            block_item.cell(row+1, col).text = block_item.cell(row+1, col).text + QuestionUI 
                                        else: 
                                            Insert_Text_Cell (block_item.cell(row, col), '' , '\n' + QuestionUI)
                                    elif len (block_item.columns) > 1 :
                                        if len (block_item.columns) > col+1 :
                                            if block_item.cell(row, col+1).text.strip() == '': 
                                                block_item.cell(row, col+1).text = block_item.cell(row, col+1).text + QuestionUI 
                                            else: 
                                                Insert_Text_Cell (block_item.cell(row, col), '' , '\n' + QuestionUI)
                                        else: 
                                            Insert_Text_Cell (block_item.cell(row, col), '' , '\n' + QuestionUI)

                                    new_dict = DictQuestions.copy() 
                                    ListDict.append ( new_dict ) 
                                    DictQuestions.clear() 
            print(ListDict)
            

            # =============== modifs pour streamlit
            # docWork.save(os.path.join(PathForOutputsAndLogs, NameOfWorkDocument + '-with UID.docx'))
            output_path = PathForOutputsAndLogs / f"{NameOfWorkDocument}-with UID.docx"

            # Sauvegarde du document
            docWork.save(output_path)
            

            # === Effacement des fichiers sources
            # try: 
            #     for file in glob.glob(PathFolderSource +'*.*'):
            #         os.remove(file)
            # except FileNotFoundError: 
            #     MessageError = str(datetime.now()) + ' Error encountered when deleting source docx file at the end of read_AAP function, please check' 
            #     logging.error(MessageError)
            #     print(MessageError)


            # Suppression des fichiers avec PathFolderSource existant
            for file_path in PathFolderSource.glob('*.*'):
                try:
                    file_path.unlink()  # Méthode pathlib pour supprimer
                    print(f"Supprimé : {file_path}")
                except Exception as e:
                    print(f"Erreur lors de la suppression de {file_path} : {e}")

            #====================== fin modifs streamlit

        else:
            MessageError = str(datetime.now()) + ' Error encountered when reading Word docx file , please check type .docx and name of the file with no UID)' 
            logging.error(MessageError)
            print(MessageError)

    print('End of the read program')
    return ListDict


# Fonction pour écrire les réponses dans l aap et dans un doc Q&A
def Write_Answers_in_docx(List_UIDQuestionsSizeAnswer, PathFolderSource, PathForOutputsAndLogs, TagQStart="<>", TagQEnd="</>"):
    """
    Remplace les UIDs dans les fichiers Word traités par les réponses correspondantes.
    Génère un fichier Word annoté + un fichier récapitulatif Q/A.
    """

    for file in glob.glob(os.path.join(PathFolderSource, '*.docx')):
        if "UID" not in os.path.basename(file):
            continue  # On ne traite que les fichiers avec UID

        try:
            with open(file, 'rb') as f:
                document = Document(f)

            NameOfDocument = os.path.basename(file)

            # === Nettoyage des balises dans les paragraphes
            for para in document.paragraphs:
                for marker in ["??", TagQStart, TagQEnd]:
                    if marker in para.text:
                        Delete_Text_Paragraph(para, marker)

            # === Remplacement des UIDs par réponses dans les paragraphes
            for para in document.paragraphs:
                for value in List_UIDQuestionsSizeAnswer:
                    if value["uid"] in para.text:
                        Insert_Text_Paragraph(para, "", "\n" + value["response"])
                        Delete_Text_Paragraph(para, value["uid"])

            # === Nettoyage et remplacement dans les tables
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for marker in ["??", TagQStart, TagQEnd]:
                            if marker in cell.text:
                                Delete_Text_Cell(cell, marker)

            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for value in List_UIDQuestionsSizeAnswer:
                            if value["uid"] in cell.text:
                                Insert_Text_Cell(cell, "", value["response"])
                                Delete_Text_Cell(cell, value["uid"])

            # === Sauvegarde du fichier final avec réponses
            timestamp = datetime.now().strftime("%Y-%m-%d_%Hh%Mmn%Ss")
            output_filename = f'{NameOfDocument.replace("-with UID", "")}_with_answers_{timestamp}.docx'
            document.save(os.path.join(PathForOutputsAndLogs, output_filename))

            # === Génération du document Q&A séparé
            documentQA = Document()
            documentQA.add_heading(
                f'Liste des questions/réponses pour :\n{NameOfDocument}\nHeure : {timestamp}\n', level=1
            )

            for value in List_UIDQuestionsSizeAnswer:
                p = documentQA.add_paragraph()
                question_clean = value["question"].replace("??", "").replace(TagQStart, "").replace(TagQEnd, "")
                run = p.add_run(question_clean)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)
                documentQA.add_paragraph('\n' + value["response"] + '\n')

            qa_filename = f'{NameOfDocument.replace("-with UID", "")}_Q-A_{timestamp}.docx'
            path_output_doc = os.path.join(PathForOutputsAndLogs, output_filename)
            path_qa_doc = os.path.join(PathForOutputsAndLogs, qa_filename)

            documentQA.save(path_qa_doc)
            documentQA = Document()

            # === Effacement des fichiers sources "avec UID"
            try: 
                for file in glob.glob(PathFolderSource +'*.*'):
                    if "UID" in os.path.basename(file):
                        os.remove(file)
            except FileNotFoundError: 
                MessageError = str(datetime.now()) + ' Error encountered when deleting source docx file at the end of write_AAP function, please check' 
                logging.error(MessageError)
                print(MessageError)

            return path_output_doc, path_qa_doc

        except Exception as e:
            print(f"[ERREUR] Problème avec le fichier {file} : {e}")

    print('✅ Fin du programme d’écriture des réponses dans les fichiers.')


# Fonctions à lancer dans le main 

## Définition des arguments de la fonction Read_Questions
Path_where_we_put_Outputs = r'./LOG'
Folder_where_the_files_are = r'./AAP'

list_of_SizeWords_OK = [
     " MAX", " MIN", " CARACT", " CHARACT", " LIGNE", " LINE", " SIGN", " PAGE",  " PAS EXC", " NOT EXCEED", " MOTS", " WORDS"
         ]

list_of_SizeWords_KO = [
     " SIGNAT", " MAXIMI", " MONTH", " MOIS", " ANS", " ANNé", " YEAR",  " DAY", " JOUR",
     " DURéE", " DURATION", " IMPACT", " AMOUNT", " MONTANT"
         ]

TagQStart = "<>"
TagQEnd = "</>"

# logging.basicConfig(filename=Path_where_we_put_Outputs + r'/logs-IA_for_Asso.txt')


