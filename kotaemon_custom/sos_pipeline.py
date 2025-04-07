

"""
sos_pipeline.py

Module gérant l’ingestion des documents, l’extraction de texte et des questions,
ainsi que la génération de réponses à partir du contexte.
"""

import os
import glob
import PyPDF2
import docx
import shutil
import re
from io import BytesIO
import gradio as gr

from kotaemon.base import Document
from kotaemon.storages import LanceDBDocumentStore
from kotaemon.storages.vectorstores.chroma import ChromaVectorStore
from kotaemon.embeddings import OpenAIEmbeddings
from llama_index.core.vector_stores.types import VectorStoreQuery

from sos_flowsettings import KH_CHAT_LLM, KH_VECTORSTORE

# Chemins pour le docstore et le vector store
DOCSTORE_PATH = "./ktem_app_data/user_data/docstore"
CHROMA_PATH   = "./ktem_app_data/user_data/chroma"

# Récupère DATA_ROOT (monté dans /app/data)
DATA_ROOT = os.environ.get("DATA_ROOT", "/app/data")
print(f"[sos_pipeline] DATA_ROOT={DATA_ROOT}")

# Mode d'embeddings : "openai", "local", ou "none"
EMBEDDING_MODE = os.environ.get("EMBEDDING_MODE", "local").lower()
print(f"[sos_pipeline] EMBEDDING_MODE={EMBEDDING_MODE}")

# Wrapper pour HuggingFaceEmbeddings via langchain
class LangchainHuggingFaceEmbeddings:
    def __init__(self, model_name="sentence-transformers/all-MiniLM-L6-v2"):
        from langchain_community.embeddings import HuggingFaceEmbeddings
        self.hf = HuggingFaceEmbeddings(model_name=model_name)

    def __call__(self, texts):
        embeddings = self.hf.embed_documents(texts)
        return [emb.tolist() if hasattr(emb, "tolist") else emb for emb in embeddings]

def read_pdf_text(pdf_path: str) -> str:
    """
    Lit et renvoie le texte d'un fichier PDF.
    """
    text = ""
    with open(pdf_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
    return text

def read_docx_text(docx_path: str) -> str:
    """
    Lit et renvoie le texte d'un fichier DOCX.
    """
    doc_file = docx.Document(docx_path)
    paragraphs = [para.text for para in doc_file.paragraphs]
    return "\n".join(paragraphs)

def main_ingest():
    """
    Parcourt le répertoire DATA_ROOT pour ingérer tous les fichiers PDF et DOCX,
    puis les ajoute au docstore et vectorstore.
    """
    docstore = LanceDBDocumentStore(path=DOCSTORE_PATH)
    vectorstore = ChromaVectorStore(
        persist_directory=CHROMA_PATH,
        collection_name="default"
    )

    # Choix du modèle d'embeddings selon le mode configuré
    if EMBEDDING_MODE == "openai":
        print("[sos_pipeline] Using OpenAIEmbeddings for ingestion.")
        embedding_model = OpenAIEmbeddings(model="text-embedding-ada-002")
    elif EMBEDDING_MODE == "local":
        print("[sos_pipeline] Using local embeddings (LangchainHuggingFaceEmbeddings) for ingestion.")
        embedding_model = LangchainHuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    else:
        print("[sos_pipeline] No embeddings will be computed (EMBEDDING_MODE != 'openai' and != 'local').")
        embedding_model = None

    # Récupération des fichiers PDF et DOCX dans DATA_ROOT
    pdf_files = glob.glob(os.path.join(DATA_ROOT, "**", "*.pdf"), recursive=True)
    docx_files = glob.glob(os.path.join(DATA_ROOT, "**", "*.docx"), recursive=True)
    all_files = pdf_files + docx_files

    for file_path in all_files:
        print(f"Ingesting: {file_path}")
        if file_path.lower().endswith(".pdf"):
            text_content = read_pdf_text(file_path)
        elif file_path.lower().endswith(".docx"):
            text_content = read_docx_text(file_path)
        else:
            print(f"[sos_pipeline] Format non supporté: {file_path}")
            continue

        doc = Document(
            page_content=text_content,
            metadata={"source": os.path.basename(file_path)}
        )
        doc_ids = docstore.add([doc])
        doc_id = doc_ids[0] if doc_ids else None

        if embedding_model is not None:
            embedding_list = embedding_model([text_content])
            embedding = embedding_list[0]
            vectorstore.add(
                [embedding],
                metadatas=[{"doc_id": doc_id, "source": os.path.basename(file_path), "text": text_content}]
            )

    print("[sos_pipeline] Ingestion complete!")

def reset_data():
    """
    Supprime les répertoires du docstore et du vectorstore.
    """
    try:
        shutil.rmtree(DOCSTORE_PATH)
        shutil.rmtree(CHROMA_PATH)
        return "Docstore et vectorstore supprimés."
    except Exception as e:
        return f"Erreur : {str(e)}"

def extract_questions_from_aap(file):
    """
    Extrait les questions d'un fichier AAP (.docx ou .pdf).

    La fonction lit le contenu du fichier, puis parcourt les lignes pour extraire
    celles qui ressemblent à des questions (en se basant sur la ponctuation ou
    certains mots-clés).
    """
    if file is None:
        return "Aucun fichier fourni."
    file_path = file.name if hasattr(file, "name") else str(file)
    ext = file_path.lower().split('.')[-1]
    if ext not in ("docx", "pdf"):
        return "Format de fichier non supporté. Veuillez fournir un .docx ou .pdf."

    text_content = ""
    try:
        if ext == "docx":
            from docx import Document
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs]
            text_content = "\n".join(paragraphs)
        elif ext == "pdf":
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            for page in doc:
                text_content += page.get_text()
            doc.close()
    except Exception as e:
        return f"Erreur lors de la lecture du fichier : {str(e)}"

    if not text_content.strip():
        return "Le document ne contient pas de texte exploitable."

    questions = []
    interrogative_words = ["qui", "quoi", "où", "quand", "comment", "pourquoi", "quel", "quelle", "quels", "quelles"]
    imperative_keywords = ["indiquez", "décrivez", "expliquez", "fournissez", "présentez", "veuillez"]

    for line in text_content.splitlines():
        if not line.strip():
            continue
        line_stripped = line.strip()
        if line_stripped.endswith("?"):
            questions.append(line_stripped)
        elif "?" in line_stripped:
            parts = line_stripped.split("?")
            for part in parts[:-1]:
                question = part.strip() + "?"
                if question:
                    questions.append(question)
        else:
            lower_line = line_stripped.lower()
            main_text = re.sub(r'^[\d\.\)\-*\s]+', '', lower_line)
            first_word = main_text.split()[0] if main_text else ""
            if first_word in interrogative_words or first_word in imperative_keywords:
                questions.append(line_stripped)

    # Suppression des doublons en conservant l'ordre
    seen = set()
    unique_questions = []
    for q in questions:
        if q not in seen:
            seen.add(q)
            unique_questions.append(q)

    if not unique_questions:
        return "Aucune question n'a été trouvée dans le document."
    return "\n".join(unique_questions)

def assistant_aap_ui():
    """
    Interface Gradio pour téléverser un AAP et afficher les questions extraites.
    """
    with gr.Blocks() as demo:
        aap_file_input = gr.File(label="Téléverser un AAP (docx/pdf)", file_types=[".docx", ".pdf"])
        questions_output = gr.Textbox(
            label="Questions extraites",
            lines=10,
            placeholder="Les questions extraites apparaîtront ici...",
            interactive=True
        )
        def on_file_uploaded(file):
            return extract_questions_from_aap(file)
        aap_file_input.change(fn=on_file_uploaded, inputs=aap_file_input, outputs=questions_output)
    return demo

def answer_question_with_context(question: str) -> str:
    """
    Génère une réponse en se basant sur le contexte extrait des documents.
    """
    print(f"[sos_pipeline] Question reçue : {question}")
    query_obj = VectorStoreQuery(
        query_str=question,
        similarity_top_k=3
    )
    results = KH_VECTORSTORE._client.query(query_obj)
    if not results or not results.nodes:
        return "Aucun contexte trouvé pour répondre à la question."

    context = "\n\n".join([node.text for node in results.nodes])
    prompt = f"""Réponds à la question suivante en te basant uniquement sur le contexte fourni.

### Contexte :
{context}

### Question :
{question}

### Réponse :"""
    response = KH_CHAT_LLM.invoke(prompt)
    return response.content

if __name__ == "__main__":
    main_ingest()
